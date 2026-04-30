"""
Conference Document Formatter - Claude-powered section detection + template-based formatting.

Usage:
    python formatter.py input.docx [-o output.docx] [--dry-run]

Workflow:
    1. Extract text from input .docx
    2. Send to Claude to identify document sections
    3. Build output using pre-configured template styles
"""

import sys
import os
import re
import json
import base64
import tempfile
import subprocess
import argparse
import copy
from pathlib import Path

from dotenv import load_dotenv
load_dotenv(Path(__file__).parent / ".env")

from docx import Document
from docx.shared import Pt, Mm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree
import anthropic

from create_template import create_template


# Normalize legacy cedilla diacritics to correct comma-below Romanian characters
_DIACRITIC_MAP = str.maketrans("şţŞŢ", "șțȘȚ")

def fix_diacritics(text: str) -> str:
    """Replace cedilla s/t (ş ţ) with correct comma-below Romanian (ș ț)."""
    return text.translate(_DIACRITIC_MAP)


# Romanian prepositions/conjunctions/function words kept lowercase in title case
_RO_LOWERCASE_WORDS = {
    'și', 'si', 'de', 'a', 'în', 'in', 'pe', 'cu', 'la', 'din', 'sau', 'ori',
    'că', 'ca', 'al', 'ale', 'cel', 'cei', 'care', 'dar', 'ci', 'nici', 'fie',
    'pentru', 'prin', 'spre', 'dintre', 'printre', 'despre', 'contra',
    'ce', 'este', 'sunt', 'era', 'fi',
}

# Tokens that act as clause separators — the word after them is capitalized
_DASH_SEPARATORS = {'—', '–', '--'}


def _capitalize_word(word):
    """Capitalize the first alphabetic character; lowercase the rest."""
    for i, c in enumerate(word):
        if c.isalpha():
            return word[:i] + c.upper() + word[i + 1:].lower()
    return word


def smart_title_case_ro(text):
    """Title case that keeps Romanian prepositions/conjunctions lowercase.

    Words immediately after a dash separator (— –) are capitalized even if
    they would normally be lowercase (e.g. 'de' → 'De' after '—').

    Example: 'PRINCIPII, METRICI SI CERTIFICARE'
             -> 'Principii, Metrici si Certificare'
             '1. Introducere — de ce eficiența energetică este insuficientă'
             -> '1. Introducere — De ce Eficiența Energetică este Insuficientă'
    """
    words = text.split()
    result = []
    capitalize_next = False  # True right after a dash separator
    for i, word in enumerate(words):
        if word in _DASH_SEPARATORS:
            result.append(word)
            capitalize_next = True
            continue
        # Strip punctuation to look up the bare word
        core = re.sub(r'^[^\w]+|[^\w]+$', '', word, flags=re.UNICODE).lower()
        if i == 0 or capitalize_next or core not in _RO_LOWERCASE_WORDS:
            result.append(_capitalize_word(word))
        else:
            result.append(word.lower())
        capitalize_next = False
    return ' '.join(result)


def sentence_case_ro(text):
    """Sentence case: only the first alphabetic character is uppercase.

    Example: 'OBIECTIVE ȘI METODOLOGIE' -> 'Obiective și metodologie'
             '1.1 OBIECTIVE' -> '1.1 Obiective'
    """
    if not text:
        return text
    result = list(text.lower())
    for i, c in enumerate(result):
        if c.isalpha():
            result[i] = c.upper()
            break
    return ''.join(result)


# Section types that Claude will assign
SECTION_TYPES = [
    "title_ro",          # Romanian title (UPPERCASE)
    "author",            # Author names
    "rezumat_label",     # The word "Rezumat"
    "rezumat_text",      # Romanian abstract text
    "keywords",          # "Cuvinte cheie:" line
    "heading1",          # Major section heading (e.g., "1. Introducere")
    "heading2",          # Sub-section heading (e.g., "1.1 Obiective")
    "body",              # Regular body paragraph
    "list_item",         # Bulleted/numbered list item
    "figure_caption",    # Figure caption (e.g., "Figura 1. ...")
    "table_caption",     # Table caption (e.g., "Tabel 1. ...")
    "table_content",     # Content inside a data table (rows/cells)
    "formula",           # Mathematical formula or equation
    "formula_label",     # Bold/uppercase header label of a formula box (skipped in output)
    "formula_legend",    # "unde:" and variable definitions after formulas
    "title_en",          # English title
    "abstract_label",    # The word "Abstract"
    "abstract_text",     # English abstract text
    "bibliography_header", # "Bibliografie" header
    "bibliography_entry",  # Individual reference entry
    "empty",             # Empty/blank line
    "skip",              # Preamble/postamble/branding to remove
]

SYSTEM_PROMPT = """You are a document structure analyzer for Romanian academic/conference papers.

You will receive the text content of a .docx document, paragraph by paragraph, with indices.
Your job is to classify each paragraph into one of these section types:

- title_ro: The main Romanian title of the paper (usually UPPERCASE, centered)
- author: Author name(s) line
- rezumat_label: The heading "Rezumat" (just the label, not the text)
- rezumat_text: The Romanian abstract paragraph(s) — the actual abstract content
- keywords: "Cuvinte cheie:" line with keywords
- heading1: Major section headings (e.g., "1. INTRODUCERE", "Concluzii", etc.)
- heading2: Sub-section headings (e.g., "1.1 Obiectivele", "2.2 Stratul 2")
- body: Regular body text paragraphs
- list_item: List items (bulleted or numbered within body)
- figure_caption: Figure captions ("Figura 1.", "Fig. 2:", etc.)
- table_caption: Table captions ("Tabel 1.", "Tabelul 2.", etc.)
- table_content: Content inside a data table (rows/cells) — marked with [TABLE ...] prefix
- formula: Mathematical formulas or equations (standalone)
- formula_label: Bold/uppercase title at the top of a formula box (e.g. "BILANȚ ENERGETIC — SISTEM DESCHIS (Legea I)") — these repeat the subchapter heading and should be removed
- formula_legend: Variable definitions following formulas ("unde:" and the definitions)
- title_en: English title of the paper
- abstract_label: The heading "Abstract" (English, just the label)
- abstract_text: English abstract paragraph(s)
- bibliography_header: The "Bibliografie" or "Referințe bibliografice" heading
- bibliography_entry: Individual bibliography/reference entries
- empty: Empty lines (preserve where structurally meaningful)
- skip: Content to remove (logos, branding, version info, organization headers, footers, "LUCRARE ȘTIINȚIFICĂ" labels)

Rules:
1. The paper structure typically follows this order: title_ro → authors → rezumat (+ cuvinte cheie) → body sections → title_en → abstract_en → bibliografie
2. Headings often have numbering (1., 2., 1.1, etc.) but not always
3. "Rezumat" and "Abstract" are labels, not body text
4. Bibliography entries usually start with [nr.] or are formatted as author-year references
5. Skip organizational branding, logos text, version numbers, and any preamble/postamble that isn't part of the paper itself
5b. Skip sub-chapter headings within the bibliography section (e.g., "10.1 Termodinamică...", "10.2 Sisteme...") — only keep the main bibliography header and the actual entries
6. Empty paragraphs between sections should be marked as "empty"
7. Figure/table captions are short, start with "Figura"/"Fig."/"Tabel"/"Tabelul", and are typically centered
8. Formula legends start with "unde:" followed by variable definitions
9. Items marked with [TABLE] prefix are table content — classify as table_content unless they clearly belong to another type (e.g., a single-cell table containing the abstract should be rezumat_text)
10. The abstract text may be inside a table cell — still classify it as rezumat_text or abstract_text based on language
11. Paragraphs marked [in-single-cell-table] that contain equations or math expressions (with =, +, –, variables, dots) should be classified as formula; lines starting with "unde:" following a formula should be formula_legend; bold/uppercase descriptive labels at the top of formula boxes (e.g. "BILANȚ ENERGETIC — SISTEM DESCHIS (Legea I)") should be classified as formula_label
11. "Cuvinte cheie:" lines contain keywords and should be classified as keywords

Respond with a JSON array where each element is:
{"idx": paragraph_index, "type": "section_type"}

Only output the JSON array, no other text. Include ALL paragraphs."""


def extract_paragraphs(doc_path):
    """Extract all content from a .docx file in document order.

    Walks through the body XML to get paragraphs AND table content
    in the correct order (not just doc.paragraphs which skips tables).
    """
    doc = Document(doc_path)
    elements = []
    idx = 0

    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    tag_p = f'{{{ns_w}}}p'
    tag_tbl = f'{{{ns_w}}}tbl'
    tag_tc = f'{{{ns_w}}}tc'
    tag_tr = f'{{{ns_w}}}tr'
    tag_t = f'{{{ns_w}}}t'
    tag_rpr = f'{{{ns_w}}}rPr'
    tag_b = f'{{{ns_w}}}b'
    tag_i = f'{{{ns_w}}}i'
    tag_jc = f'{{{ns_w}}}jc'
    tag_pstyle = f'{{{ns_w}}}pStyle'
    tag_ppr = f'{{{ns_w}}}pPr'
    tag_numPr = f'{{{ns_w}}}numPr'
    tag_numId = f'{{{ns_w}}}numId'
    tag_ilvl = f'{{{ns_w}}}ilvl'
    tag_ind = f'{{{ns_w}}}ind'

    # --- Parse numbering definitions to resolve auto-generated numbers/bullets ---
    numbering_map = {}  # numId -> {abstractNumId, levels: {ilvl: {fmt, text, left, hanging}}}
    num_counters = {}   # (numId, ilvl) -> current counter value
    try:
        numbering_part = doc.part.numbering_part
        if numbering_part:
            num_xml = numbering_part._element
            # Parse abstract numbering definitions
            abstract_nums = {}
            for absNum in num_xml.findall(f'{{{ns_w}}}abstractNum'):
                abs_id = absNum.get(f'{{{ns_w}}}abstractNumId')
                levels = {}
                for lvl in absNum.findall(f'{{{ns_w}}}lvl'):
                    ilvl_val = lvl.get(f'{{{ns_w}}}ilvl')
                    numFmt = lvl.find(f'{{{ns_w}}}numFmt')
                    lvlText = lvl.find(f'{{{ns_w}}}lvlText')
                    pPr = lvl.find(f'{{{ns_w}}}pPr')
                    ind = pPr.find(f'{{{ns_w}}}ind') if pPr is not None else None
                    levels[ilvl_val] = {
                        'fmt': numFmt.get(f'{{{ns_w}}}val') if numFmt is not None else None,
                        'text': lvlText.get(f'{{{ns_w}}}val') if lvlText is not None else None,
                        'left': ind.get(f'{{{ns_w}}}left') if ind is not None else None,
                        'hanging': ind.get(f'{{{ns_w}}}hanging') if ind is not None else None,
                    }
                abstract_nums[abs_id] = levels

            # Map numId -> abstractNumId
            for num_elem in num_xml.findall(f'{{{ns_w}}}num'):
                nid = num_elem.get(f'{{{ns_w}}}numId')
                abs_ref = num_elem.find(f'{{{ns_w}}}abstractNumId')
                abs_id = abs_ref.get(f'{{{ns_w}}}val') if abs_ref is not None else None
                if abs_id and abs_id in abstract_nums:
                    numbering_map[nid] = abstract_nums[abs_id]
    except Exception:
        pass  # No numbering part or parse error — proceed without

    def resolve_numbering(p_elem):
        """Check if paragraph has numbering and return the prefix string + indent info."""
        ppr = p_elem.find(tag_ppr)
        if ppr is None:
            return None, None, None
        numPr = ppr.find(tag_numPr)
        if numPr is None:
            return None, None, None
        numId_elem = numPr.find(tag_numId)
        ilvl_elem = numPr.find(tag_ilvl)
        nid = numId_elem.get(f'{{{ns_w}}}val') if numId_elem is not None else None
        ilvl = ilvl_elem.get(f'{{{ns_w}}}val', '0') if ilvl_elem is not None else '0'

        if nid is None or nid == '0' or nid not in numbering_map:
            return None, None, None

        lvl_info = numbering_map[nid].get(ilvl, {})
        fmt = lvl_info.get('fmt')
        text_template = lvl_info.get('text', '')
        left_emu = int(lvl_info.get('left', 0)) if lvl_info.get('left') else None
        hanging_emu = int(lvl_info.get('hanging', 0)) if lvl_info.get('hanging') else None

        prefix = ''
        if fmt == 'bullet':
            # Use the bullet character from the template
            bullet_char = text_template if text_template else '•'
            prefix = bullet_char + ' '
        elif fmt == 'decimal':
            counter_key = (nid, ilvl)
            num_counters[counter_key] = num_counters.get(counter_key, 0) + 1
            # Replace %1 with counter value
            num_text = text_template.replace('%1', str(num_counters[counter_key])) if text_template else str(num_counters[counter_key]) + '.'
            prefix = num_text + ' '
        elif fmt == 'lowerLetter':
            counter_key = (nid, ilvl)
            num_counters[counter_key] = num_counters.get(counter_key, 0) + 1
            letter = chr(ord('a') + num_counters[counter_key] - 1)
            num_text = text_template.replace('%1', letter) if text_template else letter + ')'
            prefix = num_text + ' '
        elif fmt == 'upperLetter':
            counter_key = (nid, ilvl)
            num_counters[counter_key] = num_counters.get(counter_key, 0) + 1
            letter = chr(ord('A') + num_counters[counter_key] - 1)
            num_text = text_template.replace('%1', letter) if text_template else letter + ')'
            prefix = num_text + ' '

        return prefix, left_emu, hanging_emu

    def get_para_info(p_elem):
        """Extract text and metadata from a paragraph XML element."""
        text_parts = []
        is_bold = False
        is_italic = False
        for r in p_elem:
            if r.tag == tag_rpr:
                continue
            # Check run properties
            rpr = r.find(tag_rpr)
            if rpr is not None:
                b_elem = rpr.find(tag_b)
                if b_elem is not None and b_elem.get(f'{{{ns_w}}}val', 'true') != 'false':
                    is_bold = True
                i_elem = rpr.find(tag_i)
                if i_elem is not None and i_elem.get(f'{{{ns_w}}}val', 'true') != 'false':
                    is_italic = True
            for t_elem in r.iter(tag_t):
                if t_elem.text:
                    text_parts.append(t_elem.text)

        text = fix_diacritics(''.join(text_parts).strip())

        # Get paragraph properties
        alignment = "None"
        style = "Normal"
        ppr = p_elem.find(tag_ppr)
        if ppr is not None:
            jc = ppr.find(tag_jc)
            if jc is not None:
                alignment = jc.get(f'{{{ns_w}}}val', 'None')
            ps = ppr.find(tag_pstyle)
            if ps is not None:
                style = ps.get(f'{{{ns_w}}}val', 'Normal')

        return text, style, alignment, is_bold, is_italic

    def process_table(tbl_elem):
        """Extract table content as a series of elements."""
        nonlocal idx
        rows = tbl_elem.findall(tag_tr)
        num_rows = len(rows)
        num_cols = 0
        if rows:
            num_cols = len(rows[0].findall(tag_tc))

        # Single-column tables (1 col, any rows) are used for formula boxes and
        # abstract boxes — extract their paragraphs inline so Claude can classify
        # each row individually (formula, formula_legend, rezumat_text, etc.)
        is_single_col = (num_cols == 1)

        if is_single_col:
            # Extract paragraphs from each row's single cell directly
            for row in rows:
                cell = row.find(tag_tc)
                if cell is None:
                    continue
                for p_elem in cell.findall(tag_p):
                    text, style, alignment, is_bold, is_italic = get_para_info(p_elem)
                    elements.append({
                        "idx": idx,
                        "text": text,
                        "style": style,
                        "alignment": alignment,
                        "is_bold": is_bold,
                        "is_italic": is_italic,
                        "is_empty": len(text) == 0,
                        "is_table_cell": True,
                        "table_type": "single_cell",
                    })
                    idx += 1
        else:
            # Multi-cell data table — store as a single element with grid data
            table_data = []
            for ri, row in enumerate(rows):
                cells = row.findall(tag_tc)
                row_data = []
                for cell in cells:
                    cell_text_parts = []
                    for p_elem in cell.findall(tag_p):
                        text, _, _, _, _ = get_para_info(p_elem)
                        if text:
                            cell_text_parts.append(text)
                    row_data.append('\n'.join(cell_text_parts))
                table_data.append(row_data)

            # Detect split tables: if the previous element is a data table with the
            # same number of columns AND the same first row (repeated page-break header),
            # merge the rows instead of creating a separate table element.
            prev = elements[-1] if elements else None
            if (prev and
                    prev.get("table_type") == "data_table" and
                    prev.get("table_cols") == num_cols and
                    table_data and prev["table_data"] and
                    table_data[0] == prev["table_data"][0]):
                # Continuation table — append non-header rows to the previous element
                prev["table_data"].extend(table_data[1:])
                prev["table_rows"] = len(prev["table_data"])
                header_text = ' | '.join(prev["table_data"][0])
                summary = f"[TABLE {prev['table_rows']}x{num_cols}] Header: {header_text}"
                if len(summary) > 200:
                    summary = summary[:200] + "..."
                prev["text"] = summary
                # Don't create a new element — no idx increment
            else:
                header_text = ' | '.join(table_data[0]) if table_data else ''
                summary = f"[TABLE {num_rows}x{num_cols}] Header: {header_text}"
                if len(summary) > 200:
                    summary = summary[:200] + "..."

                elements.append({
                    "idx": idx,
                    "text": summary,
                    "style": "Table",
                    "alignment": "None",
                    "is_bold": False,
                    "is_italic": False,
                    "is_empty": False,
                    "is_table_cell": True,
                    "table_type": "data_table",
                    "table_rows": num_rows,
                    "table_cols": num_cols,
                    "table_data": table_data,
                })
                idx += 1

    # Walk through body children in order
    body = doc.element.body
    for child in body:
        if child.tag == tag_p:
            text, style, alignment, is_bold, is_italic = get_para_info(child)

            # Resolve auto-numbering and prepend prefix
            list_prefix, list_left, list_hanging = resolve_numbering(child)
            if list_prefix and text:
                text = list_prefix + text

            # Capture explicit indent from source paragraph
            src_left_indent = None
            src_first_indent = None
            ppr = child.find(tag_ppr)
            if ppr is not None:
                ind = ppr.find(tag_ind)
                if ind is not None:
                    left_val = ind.get(f'{{{ns_w}}}left')
                    hanging_val = ind.get(f'{{{ns_w}}}hanging')
                    first_val = ind.get(f'{{{ns_w}}}firstLine')
                    if left_val:
                        src_left_indent = int(left_val)
                    if hanging_val:
                        src_first_indent = -int(hanging_val)
                    elif first_val:
                        src_first_indent = int(first_val)
            # If no explicit indent but numbering defines one, use numbering indent
            if list_left is not None and src_left_indent is None:
                src_left_indent = list_left
            if list_hanging is not None and src_first_indent is None:
                src_first_indent = -list_hanging

            elements.append({
                "idx": idx,
                "text": text,
                "style": style,
                "alignment": alignment,
                "is_bold": is_bold,
                "is_italic": is_italic,
                "is_empty": len(text) == 0,
                "is_table_cell": False,
                "table_type": None,
                "list_prefix": list_prefix,
                "src_left_indent": src_left_indent,
                "src_first_indent": src_first_indent,
            })
            idx += 1
        elif child.tag == tag_tbl:
            process_table(child)

    return elements, doc


def _docx_to_pdf_bytes(docx_path: str) -> bytes:
    """Convert a DOCX to PDF bytes.

    Tries docx2pdf (Windows/Word) first, falls back to LibreOffice headless.
    """
    import glob

    with tempfile.TemporaryDirectory() as tmpdir:
        pdf_path = os.path.join(tmpdir, "doc.pdf")
        converted = False

        try:
            from docx2pdf import convert
            convert(docx_path, pdf_path)
            converted = True
        except Exception as e:
            print(f"docx2pdf unavailable ({e}), trying LibreOffice...")

        if not converted:
            result = subprocess.run(
                ["soffice", "--headless", "--convert-to", "pdf", "--outdir", tmpdir, docx_path],
                capture_output=True, text=True,
            )
            if result.returncode != 0:
                raise RuntimeError(f"LibreOffice PDF conversion failed: {result.stderr}")
            pdf_files = glob.glob(os.path.join(tmpdir, "*.pdf"))
            if not pdf_files:
                raise RuntimeError("LibreOffice did not produce a PDF output")
            pdf_path = pdf_files[0]

        with open(pdf_path, "rb") as f:
            return f.read()


def build_claude_message(paragraphs):
    """Build the message to send to Claude for section detection."""
    lines = []
    for p in paragraphs:
        meta = []
        if p["is_bold"]:
            meta.append("bold")
        if p["is_italic"]:
            meta.append("italic")
        if "center" in p["alignment"].lower():
            meta.append("centered")
        if "Heading" in p["style"]:
            meta.append(f"style:{p['style']}")
        if "Title" in p["style"]:
            meta.append("style:Title")
        if "List" in p["style"]:
            meta.append("style:ListParagraph")
        if p.get("is_table_cell"):
            if p.get("table_type") == "single_cell":
                meta.append("in-single-cell-table")
            elif p.get("table_type") == "data_table":
                meta.append(f"data-table({p.get('table_rows','?')}x{p.get('table_cols','?')})")

        meta_str = f" [{', '.join(meta)}]" if meta else ""
        text = p["text"] if p["text"] else "[EMPTY]"

        # Truncate very long text to save tokens
        if len(text) > 200:
            text = text[:200] + "..."

        lines.append(f"P{p['idx']:3d}{meta_str}: {text}")

    return "\n".join(lines)


def restore_diacritics(paragraphs, model="claude-sonnet-4-6"):
    """Use Claude to restore missing Romanian diacritics in all paragraph texts,
    including table cell contents stored in p['table_data'].

    Processes in batches of 80 entries to avoid hitting output token limits.
    """
    client = anthropic.Anthropic()

    SYSTEM = (
        "You are a Romanian text corrector. Your only job is to restore missing "
        "diacritics (ă, â, î, ș, ț and their uppercase variants Ă, Â, Î, Ș, Ț) in Romanian text. "
        "This includes ALL-CAPS Romanian text — for example 'BIOMASA' → 'BIOMASĂ', "
        "'INTEGRAREA' → 'INTEGRAREA', 'BILANT' → 'BILANȚ', 'SI' → 'ȘI', 'IN' → 'ÎN'. "
        "Do NOT change any other words, spelling, punctuation, or order. "
        "Preserve technical terms, acronyms, product names, formulas, numbers, and English words exactly. "
        "You MUST return every input line — even if unchanged. "
        "Return ONLY the corrected lines in the exact same format: key|||corrected_text. "
        "One line per input line. No extra commentary."
    )

    # Build a flat list of (key, text) for all text chunks
    entries = []
    for p in paragraphs:
        if p["text"].strip():
            entries.append((f"p{p['idx']}", p["text"]))
        if p.get("table_data"):
            for r_i, row in enumerate(p["table_data"]):
                for c_i, cell in enumerate(row):
                    if cell.strip():
                        entries.append((f"t{p['idx']}_{r_i}_{c_i}", cell))

    if not entries:
        return paragraphs

    # Process in batches to avoid output token limit
    BATCH_SIZE = 80
    batches = [entries[i:i + BATCH_SIZE] for i in range(0, len(entries), BATCH_SIZE)]
    print(f"Restoring Romanian diacritics for {len(entries)} text chunks ({len(batches)} batch(es))...")

    idx_to_para = {p["idx"]: p for p in paragraphs}
    total_in, total_out = 0, 0

    def apply_response(text):
        for line in text.strip().splitlines():
            if "|||" not in line:
                continue
            key, _, corrected = line.partition("|||")
            key = key.strip()
            corrected = corrected.strip()
            if not corrected:
                continue
            if key.startswith("p"):
                try:
                    idx = int(key[1:])
                    if idx in idx_to_para:
                        idx_to_para[idx]["text"] = corrected
                except ValueError:
                    pass
            elif key.startswith("t"):
                try:
                    parts = key[1:].split("_")
                    idx, r_i, c_i = int(parts[0]), int(parts[1]), int(parts[2])
                    if idx in idx_to_para:
                        td = idx_to_para[idx].get("table_data")
                        if td and r_i < len(td) and c_i < len(td[r_i]):
                            td[r_i][c_i] = corrected
                except (ValueError, IndexError):
                    pass

    for i, batch in enumerate(batches):
        lines = "\n".join(f"{key}|||{text}" for key, text in batch)
        response = client.messages.create(
            model=model,
            max_tokens=16384,
            system=SYSTEM,
            messages=[{"role": "user", "content": f"Restore diacritics in these Romanian text lines:\n\n{lines}"}]
        )
        total_in += response.usage.input_tokens
        total_out += response.usage.output_tokens
        sent_keys = {key for key, _ in batch}
        apply_response(response.content[0].text)
        returned_keys = {line.partition("|||")[0].strip() for line in response.content[0].text.strip().splitlines() if "|||" in line}
        missed = sent_keys - returned_keys
        if missed:
            print(f"  Batch {i+1}: {len(missed)} entries not returned by Claude: {sorted(missed)[:5]}{'...' if len(missed)>5 else ''}")

    print(f"Diacritics restoration complete. Token usage: input={total_in}, output={total_out}")
    return paragraphs


def classify_with_claude(paragraphs, model="claude-sonnet-4-6", docx_path=None):
    """Send the full document PDF + paragraph list to Claude for section classification."""
    client = anthropic.Anthropic()

    paragraph_list = build_claude_message(paragraphs)

    content = []

    if docx_path:
        print(f"Converting document to PDF for Claude vision analysis...")
        pdf_bytes = _docx_to_pdf_bytes(docx_path)
        print(f"  PDF size: {len(pdf_bytes) // 1024} KB — sending to Claude as document")
        content.append({
            "type": "document",
            "source": {
                "type": "base64",
                "media_type": "application/pdf",
                "data": base64.b64encode(pdf_bytes).decode(),
            },
        })
        prompt = (
            "The full document is attached above as a PDF. "
            "Use the visual layout — font sizes, bold/italic formatting, centering, indentation, "
            "and content — together with the paragraph list below to classify each section.\n\n"
            f"Classify each paragraph in this document:\n\n{paragraph_list}"
        )
    else:
        prompt = f"Classify each paragraph in this document:\n\n{paragraph_list}"

    content.append({"type": "text", "text": prompt})

    print(f"Sending {len(paragraphs)} paragraphs to Claude ({model})...")

    response = client.messages.create(
        model=model,
        max_tokens=16384,
        system=SYSTEM_PROMPT,
        messages=[{"role": "user", "content": content}],
    )

    # Parse response
    response_text = response.content[0].text.strip()

    # Try to extract JSON from the response
    # Sometimes Claude wraps it in ```json ... ```
    if "```" in response_text:
        match = re.search(r'```(?:json)?\s*\n?(.*?)\n?```', response_text, re.DOTALL)
        if match:
            response_text = match.group(1).strip()

    try:
        classifications = json.loads(response_text)
    except json.JSONDecodeError as e:
        print(f"Failed to parse Claude response as JSON: {e}")
        print(f"Response preview: {response_text[:500]}")
        sys.exit(1)

    # Build index map
    section_map = {}
    for item in classifications:
        section_map[item["idx"]] = item["type"]

    # Fill in any missing indices
    for p in paragraphs:
        if p["idx"] not in section_map:
            section_map[p["idx"]] = "empty" if p["is_empty"] else "body"

    print(f"Classification complete. Token usage: input={response.usage.input_tokens}, output={response.usage.output_tokens}")

    # Print summary
    type_counts = {}
    for t in section_map.values():
        type_counts[t] = type_counts.get(t, 0) + 1
    print(f"Section counts: {json.dumps(type_counts, indent=2)}")

    return section_map


def parse_bibliography_ieee(entries, model="claude-sonnet-4-6"):
    """Send bibliography entries to Claude to parse into IEEE-style structured parts.

    Returns a list of lists, where each inner list contains dicts:
      {"text": "...", "format": "bold"|"italic"|"normal"}
    """
    client = anthropic.Anthropic()

    entries_text = "\n".join(f"[{i}] {e}" for i, e in enumerate(entries))

    print(f"Parsing {len(entries)} bibliography entries for IEEE formatting...")

    response = client.messages.create(
        model=model,
        max_tokens=8192,
        system="""You are a bibliography formatter. You will receive a list of bibliography entries.
For each entry, split it into three parts following IEEE style:
1. "number_and_authors" — the reference number (e.g. [1]) and author names → bold
2. "title" — the title of the paper/book/chapter → italic
3. "rest" — everything else (journal, volume, pages, year, DOI, URL, etc.) → normal

Respond with a JSON array where each element corresponds to an entry:
[
  {
    "idx": 0,
    "parts": [
      {"text": "[1] Author Name, Another Author", "format": "bold"},
      {"text": ", \"Title of the Paper\",", "format": "italic"},
      {"text": " Journal Name, vol. 1, pp. 1-10, 2020.", "format": "normal"}
    ]
  },
  ...
]

Rules:
- Keep punctuation attached to the part it belongs to (commas, periods, quotes)
- The number prefix like [1], [2] etc. goes with the authors (bold part)
- If you cannot identify a clear title, make the entire entry bold (as-is)
- Preserve all original text exactly — do not rewrite or reorder
- Only output the JSON array, no other text""",
        messages=[
            {
                "role": "user",
                "content": f"Parse these bibliography entries:\n\n{entries_text}"
            }
        ]
    )

    response_text = response.content[0].text.strip()
    if "```" in response_text:
        match = re.search(r'```(?:json)?\s*\n?(.*?)\n?```', response_text, re.DOTALL)
        if match:
            response_text = match.group(1).strip()

    try:
        parsed = json.loads(response_text)
    except json.JSONDecodeError as e:
        print(f"Failed to parse bibliography response: {e}")
        # Fallback: return all entries as bold (original behavior)
        return [[{"text": e, "format": "bold"}] for e in entries]

    # Build result indexed by position
    result = []
    parsed_map = {item["idx"]: item["parts"] for item in parsed}
    for i in range(len(entries)):
        if i in parsed_map:
            result.append(parsed_map[i])
        else:
            result.append([{"text": entries[i], "format": "bold"}])

    print(f"Bibliography parsing complete. Token usage: input={response.usage.input_tokens}, output={response.usage.output_tokens}")
    return result


def assign_figure_table_numbers(paragraphs, section_map):
    """Pre-pass to assign sequential numbers to all figures, tables, and formulas.

    Numbers are sequential across the whole document (1, 2, 3, ...).

    Handles both caption-before-table and caption-after-table patterns:
    both elements in a pair share the same number.

    Returns:
        figure_numbers:   dict mapping paragraph idx -> number string
        table_numbers:    dict mapping paragraph idx -> number string
        formula_numbers:  dict mapping paragraph idx -> number string
        post_caption_map: dict mapping table_content_idx -> table_caption_idx
                          for caption-after-table pairs (so build loop can use
                          the explicit caption text and skip the post-caption later)
    """
    # --- Pass 1: detect adjacent (table_content, table_caption) pairs in both orders ---
    pre_caption_content = {}   # table_caption_idx -> table_content_idx  (caption BEFORE table)
    post_caption_map = {}      # table_content_idx -> table_caption_idx  (caption AFTER table)

    prev_sig_type = None
    prev_sig_idx = None
    for p_info in paragraphs:
        sec_type = section_map.get(p_info["idx"], "body")
        if sec_type in ("skip", "empty"):
            continue
        cur_idx = p_info["idx"]

        if sec_type == "table_caption" and prev_sig_type == "table_content":
            # Caption comes right after the table data
            post_caption_map[prev_sig_idx] = cur_idx
        elif sec_type == "table_content" and p_info.get("table_data") and prev_sig_type == "table_caption":
            # Caption came right before the table data
            pre_caption_content[prev_sig_idx] = cur_idx

        prev_sig_type = sec_type
        prev_sig_idx = cur_idx

    # --- Pass 2: assign sequential numbers, counting each (content+caption) pair as one table ---
    figure_numbers = {}
    table_numbers = {}
    formula_numbers = {}
    fig_counter = 0
    tbl_counter = 0
    formula_counter = 0
    already_numbered = set()  # idx of elements already given a number via pairing

    for p_info in paragraphs:
        sec_type = section_map.get(p_info["idx"], "body")
        if sec_type in ("skip", "empty"):
            continue
        cur_idx = p_info["idx"]

        if sec_type == "figure_caption":
            if cur_idx in already_numbered:
                continue
            fig_counter += 1
            num_str = str(fig_counter)
            figure_numbers[cur_idx] = num_str

        elif sec_type == "table_caption":
            if cur_idx in already_numbered:
                continue
            tbl_counter += 1
            num_str = str(tbl_counter)
            table_numbers[cur_idx] = num_str
            # If this caption is paired with a following table_content, give it the same number
            paired_content = pre_caption_content.get(cur_idx)
            if paired_content is not None:
                table_numbers[paired_content] = num_str
                already_numbered.add(paired_content)

        elif sec_type == "table_content" and p_info.get("table_data"):
            if cur_idx in already_numbered:
                continue
            tbl_counter += 1
            num_str = str(tbl_counter)
            table_numbers[cur_idx] = num_str
            # If this content is paired with a following table_caption, give it the same number
            paired_caption = post_caption_map.get(cur_idx)
            if paired_caption is not None:
                table_numbers[paired_caption] = num_str
                already_numbered.add(paired_caption)

        elif sec_type == "formula":
            if cur_idx in already_numbered:
                continue
            formula_counter += 1
            formula_numbers[cur_idx] = str(formula_counter)

    return figure_numbers, table_numbers, formula_numbers, post_caption_map


def find_existing_references(paragraphs, section_map):
    """Scan body text for existing figure/table references.

    Returns:
        referenced_figures: set of figure numbers found in body text
        referenced_tables: set of table numbers found in body text
    """
    fig_pattern = re.compile(r'(?:Figura|Fig\.?)\s*(\d+(?:\.\d+)?)', re.IGNORECASE)
    tbl_pattern = re.compile(r'(?:Tabel(?:ul)?)\s*(\d+(?:\.\d+)?)', re.IGNORECASE)
    scannable_types = {"body", "list_item", "formula_legend", "rezumat_text"}

    referenced_figures = set()
    referenced_tables = set()
    for p_info in paragraphs:
        sec_type = section_map.get(p_info["idx"], "body")
        if sec_type not in scannable_types:
            continue
        text = p_info.get("text", "")
        for m in fig_pattern.finditer(text):
            referenced_figures.add(m.group(1))
        for m in tbl_pattern.finditer(text):
            referenced_tables.add(m.group(1))
    return referenced_figures, referenced_tables


def generate_reference_sentences(unreferenced_items, model="claude-sonnet-4-6"):
    """Generate natural Romanian introductory sentences for unreferenced tables/figures.

    Args:
        unreferenced_items: list of dicts with keys: number, item_type ("table"/"figure"), caption
    Returns:
        dict mapping (item_type, number) -> generated sentence
    """
    if not unreferenced_items:
        return {}

    client = anthropic.Anthropic()

    items_text = "\n".join(
        f"{item['item_type'].capitalize()} {item['number']}: {item['caption']}"
        for item in unreferenced_items
    )

    print(f"Generating reference sentences for {len(unreferenced_items)} unreferenced items...")

    response = client.messages.create(
        model=model,
        max_tokens=4096,
        system="""You generate short Romanian introductory sentences for tables and figures in academic papers.
Each sentence should naturally introduce the table/figure using its caption as context.

Examples:
- For "Table 8: Structura pe niveluri cu obiective și parametri financiari":
  "Structura pe niveluri cu obiective și parametri financiari se poate observa în Tabelul 8."
- For "Table 3: Componentele din perspectiva entropiei":
  "Componentele analizate din perspectiva entropiei sunt prezentate în Tabelul 3."
- For "Figure 1: Schema ierarhiei AHP":
  "Schema ierarhiei AHP este ilustrată în Figura 1."

Rules:
- Write in Romanian
- Each sentence should be 1 sentence, concise (under 20 words)
- Reference the table/figure naturally using "Tabelul X" or "Figura X"
- Use varied phrasing: "se poate observa în", "este prezentat/ă în", "sunt prezentate în", "este ilustrat/ă în", etc.
- Only output JSON, no other text

Respond with a JSON array:
[{"type": "table", "number": 8, "sentence": "..."}, ...]""",
        messages=[
            {
                "role": "user",
                "content": f"Generate introductory sentences for:\n\n{items_text}"
            }
        ]
    )

    response_text = response.content[0].text.strip()
    if "```" in response_text:
        match = re.search(r'```(?:json)?\s*\n?(.*?)\n?```', response_text, re.DOTALL)
        if match:
            response_text = match.group(1).strip()

    result = {}
    try:
        parsed = json.loads(response_text)
        for item in parsed:
            result[(item["type"], item["number"])] = item["sentence"]
    except json.JSONDecodeError as e:
        print(f"Failed to parse reference sentences: {e}")
        # Fallback to simple pattern
        for item in unreferenced_items:
            if item["item_type"] == "table":
                result[("table", item["number"])] = f"{item['caption']} se poate observa în Tabelul {item['number']}."
            else:
                result[("figure", item["number"])] = f"{item['caption']} este ilustrată în Figura {item['number']}."

    print(f"Reference sentence generation complete. Token usage: input={response.usage.input_tokens}, output={response.usage.output_tokens}")
    return result


def build_reference_insertions(paragraphs, section_map, figure_numbers, table_numbers,
                                referenced_figures, referenced_tables, table_captions_map, model="claude-sonnet-4-6"):
    """Build a map of intro sentences to insert before unreferenced figures/tables.

    Returns:
        dict mapping paragraph idx -> intro sentence to insert before that paragraph
    """
    # Collect unreferenced items with their captions
    unreferenced_items = []
    for para_idx, fig_num in figure_numbers.items():
        if fig_num not in referenced_figures:
            # Figure captions: find the paragraph text
            caption = ""
            for p in paragraphs:
                if p["idx"] == para_idx:
                    caption = re.sub(r'^(Figura|Fig\.?)\s*\d+[\.:]\s*', '', p.get("text", ""), flags=re.IGNORECASE).strip()
                    break
            unreferenced_items.append({
                "para_idx": para_idx,
                "item_type": "figure",
                "number": fig_num,
                "caption": caption or f"Figura {fig_num}"
            })
    for para_idx, tbl_num in table_numbers.items():
        if tbl_num not in referenced_tables:
            caption = table_captions_map.get(para_idx, "")
            unreferenced_items.append({
                "para_idx": para_idx,
                "item_type": "table",
                "number": tbl_num,
                "caption": caption or f"Tabelul {tbl_num}"
            })

    if not unreferenced_items:
        return {}

    # Generate natural sentences
    sentences = generate_reference_sentences(unreferenced_items, model=model)

    # Map each unreferenced item's paragraph idx -> sentence
    insertions = {}
    for item in unreferenced_items:
        key = (item["item_type"], item["number"])
        sentence = sentences.get(key)
        if sentence:
            insertions[item["para_idx"]] = sentence

    return insertions


def generate_table_captions(table_infos, model="claude-sonnet-4-6"):
    """Generate descriptive Romanian captions for tables based on their header rows.

    Args:
        table_infos: list of (paragraph_idx, header_summary) tuples
    Returns:
        dict mapping paragraph_idx -> caption text (without "Tabel X." prefix)
    """
    client = anthropic.Anthropic()

    tables_text = "\n".join(f"Table {i}: columns = {header}" for i, (_, header) in enumerate(table_infos))

    print(f"Generating captions for {len(table_infos)} tables...")

    response = client.messages.create(
        model=model,
        max_tokens=4096,
        system="""You generate short, descriptive Romanian captions for tables in academic papers.
Given the column headers of each table, produce a concise caption that describes what the table contains.

Respond with a JSON array:
[{"idx": 0, "caption": "Descrierea concisă a tabelului"}, ...]

Rules:
- Captions should be in Romanian
- Keep them concise (under 15 words)
- Describe the content/purpose, not the structure
- Do not include "Tabel X." prefix — just the descriptive text
- Only output the JSON array, no other text""",
        messages=[
            {
                "role": "user",
                "content": f"Generate Romanian captions for these tables:\n\n{tables_text}"
            }
        ]
    )

    response_text = response.content[0].text.strip()
    if "```" in response_text:
        match = re.search(r'```(?:json)?\s*\n?(.*?)\n?```', response_text, re.DOTALL)
        if match:
            response_text = match.group(1).strip()

    result = {}
    try:
        parsed = json.loads(response_text)
        for item in parsed:
            table_idx = item["idx"]
            if table_idx < len(table_infos):
                para_idx = table_infos[table_idx][0]
                result[para_idx] = item["caption"]
    except json.JSONDecodeError as e:
        print(f"Failed to parse table captions response: {e}")

    print(f"Table caption generation complete. Token usage: input={response.usage.input_tokens}, output={response.usage.output_tokens}")
    return result


def translate_to_english(paragraphs_ro, model="claude-sonnet-4-6"):
    """Translate Romanian abstract paragraphs to English.

    Returns a list of translated paragraph strings.
    """
    client = anthropic.Anthropic()
    combined = "\n\n".join(paragraphs_ro)

    response = client.messages.create(
        model=model,
        max_tokens=4096,
        system="""You are a professional academic translator from Romanian to English.
Translate the given Romanian abstract text to English.
- Preserve the academic tone and technical terminology
- Keep the same paragraph structure (one translated paragraph per input paragraph)
- Output only the translated text, no explanations
- If the input has multiple paragraphs separated by blank lines, preserve that separation""",
        messages=[{"role": "user", "content": f"Translate this Romanian abstract to English:\n\n{combined}"}]
    )

    translated = response.content[0].text.strip()
    print(f"Translation complete. Token usage: input={response.usage.input_tokens}, output={response.usage.output_tokens}")

    # Split back into paragraphs if there were multiple
    parts = [p.strip() for p in translated.split("\n\n") if p.strip()]
    return parts if parts else [translated]


def build_formatted_document(doc_path, section_map, paragraphs, template_path, model="claude-sonnet-4-6", authors=None, title_en=None, fast=False):
    """Build the formatted output document using the template styles."""

    # Reuse existing template if present, otherwise create it once
    if not os.path.exists(template_path):
        create_template(template_path)
    out_doc = Document(template_path)

    # Remove any existing content
    for p in out_doc.paragraphs:
        p._element.getparent().remove(p._element)

    source_doc = Document(doc_path)
    source_paragraphs = source_doc.paragraphs

    # Collect bibliography entries and parse them for IEEE formatting
    bib_entries = []
    bib_idx_map = {}  # maps paragraph idx -> position in bib_entries list
    for p_info in paragraphs:
        if section_map.get(p_info["idx"]) == "bibliography_entry" and p_info["text"]:
            bib_idx_map[p_info["idx"]] = len(bib_entries)
            bib_entries.append(p_info["text"])

    bib_parsed = []
    if bib_entries and not fast:
        bib_parsed = parse_bibliography_ieee(bib_entries, model=model)
    elif bib_entries and fast:
        print("Skipping IEEE bibliography parsing (fast mode)")

    # Generate table captions for data tables that have no explicit caption
    table_infos = []  # list of (paragraph idx, header summary)
    for p_info in paragraphs:
        if section_map.get(p_info["idx"]) == "table_content" and p_info.get("table_data"):
            header_row = p_info["table_data"][0]
            table_infos.append((p_info["idx"], " | ".join(header_row)))

    table_captions_map = {}  # paragraph idx -> generated caption text
    if table_infos and not fast:
        table_captions_map = generate_table_captions(table_infos, model=model)
    elif table_infos and fast:
        print("Skipping table caption generation (fast mode)")

    # --- Auto-reference insertion for figures and tables ---
    figure_numbers, table_numbers, formula_numbers, post_caption_map = assign_figure_table_numbers(paragraphs, section_map)

    # Build old-number → new-number mapping from explicit source captions.
    # e.g. source caption "Tabel 7. ..." → old "7", new "4"
    old_to_new_table = {}
    old_to_new_figure = {}
    for p_info in paragraphs:
        sec_type = section_map.get(p_info["idx"])
        if sec_type == "table_caption":
            m = re.search(r'(?:Tabel(?:ul)?)\s*(\d+(?:\.\d+)?)', p_info["text"], re.IGNORECASE)
            if m:
                new_num = table_numbers.get(p_info["idx"])
                if new_num:
                    old_to_new_table[m.group(1)] = new_num
        elif sec_type == "figure_caption":
            m = re.search(r'(?:Figura|Fig\.?)\s*(\d+(?:\.\d+)?)', p_info["text"], re.IGNORECASE)
            if m:
                new_num = figure_numbers.get(p_info["idx"])
                if new_num:
                    old_to_new_figure[m.group(1)] = new_num

    # Build a fast lookup: paragraph idx -> paragraph info
    para_by_idx = {p["idx"]: p for p in paragraphs}
    # Set of table_caption idxs that are POST-captions (appear after their table data)
    post_caption_idxs = set(post_caption_map.values())

    referenced_figures, referenced_tables = find_existing_references(paragraphs, section_map)

    # Also treat tables/figures referenced by their OLD numbers as referenced,
    # so we don't generate unwanted auto-reference sentences for them.
    for old_num, new_num in old_to_new_table.items():
        if old_num in referenced_tables:
            referenced_tables.add(new_num)
    for old_num, new_num in old_to_new_figure.items():
        if old_num in referenced_figures:
            referenced_figures.add(new_num)

    ref_insertions = build_reference_insertions(
        paragraphs, section_map,
        figure_numbers, table_numbers,
        referenced_figures, referenced_tables,
        table_captions_map, model=model
    )

    def fix_table_refs(text):
        """Replace old table/figure numbers in body text with new sequential ones.

        Handles both plain integers ("Tabelul 2") and decimals ("Tabelul 2.2").
        Only replaces when a mapping is known; leaves unknown numbers untouched.
        """
        def repl_table(m):
            old = m.group(2)
            return m.group(1) + ' ' + old_to_new_table.get(old, old)
        def repl_figure(m):
            old = m.group(2)
            return m.group(1) + ' ' + old_to_new_figure.get(old, old)
        text = re.sub(r'(Tabel(?:ul)?)\s+(\d+(?:\.\d+)?)', repl_table, text, flags=re.IGNORECASE)
        text = re.sub(r'(Figura|Fig\.?)\s+(\d+(?:\.\d+)?)', repl_figure, text, flags=re.IGNORECASE)
        return text

    def add_blank_lines(n):
        for _ in range(n):
            p = out_doc.add_paragraph("", style='Normal')
            # Clear the first-line indent for blank lines
            p.paragraph_format.first_line_indent = Mm(0)

    _subscript_pattern = re.compile(r'([^\s_]+)_(\w+)')
    # Text width B5 (136mm) in twips for formula tab stops
    _TEXT_WIDTH_TWIPS = 7710

    def _add_subscript_runs(p, text):
        """Append runs with subscript formatting for base_sub patterns."""
        last_end = 0
        for m in _subscript_pattern.finditer(text):
            if m.start() > last_end:
                run = p.add_run(text[last_end:m.start()])
                run.font.name = "Times New Roman"
                run.font.size = Pt(11)
            run = p.add_run(m.group(1))
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
            run = p.add_run(m.group(2))
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
            run.font.subscript = True
            last_end = m.end()
        if last_end < len(text):
            run = p.add_run(text[last_end:])
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)

    def add_formula_paragraph(text, centered=True, formula_num=None):
        """Add a formula paragraph with subscript formatting.

        If formula_num is given, the formula is centered via a tab stop and
        the number appears right-aligned on the same line as (1), (2), etc.
        """
        p = out_doc.add_paragraph(style='Normal')
        p.paragraph_format.first_line_indent = Mm(0)

        if formula_num:
            # Set tab stops: center for formula, right for number
            pPr = p._p.get_or_add_pPr()
            tabs_xml = parse_xml(
                f'<w:tabs {nsdecls("w")}>'
                f'  <w:tab w:val="center" w:pos="{_TEXT_WIDTH_TWIPS // 2}"/>'
                f'  <w:tab w:val="right" w:pos="{_TEXT_WIDTH_TWIPS}"/>'
                f'</w:tabs>'
            )
            pPr.append(tabs_xml)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Tab → centered formula text → tab → right-aligned (number)
            run = p.add_run('\t')
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
            _add_subscript_runs(p, text)
            run = p.add_run(f'\t({formula_num})')
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
        else:
            if centered:
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_subscript_runs(p, text)

        return p

    def add_paragraph(text, style_name, **overrides):
        """Add a paragraph with the given style and optional overrides."""
        p = out_doc.add_paragraph(text, style=style_name)

        # Apply any overrides
        if 'bold' in overrides:
            for run in p.runs:
                run.font.bold = overrides['bold']
        if 'italic' in overrides:
            for run in p.runs:
                run.font.italic = overrides['italic']
        if 'font_size' in overrides:
            for run in p.runs:
                run.font.size = overrides['font_size']
        if 'alignment' in overrides:
            p.paragraph_format.alignment = overrides['alignment']
        if 'first_line_indent' in overrides:
            p.paragraph_format.first_line_indent = overrides['first_line_indent']

        return p

    # Track what we've written to manage spacing
    prev_type = None
    wrote_title = False
    wrote_authors = False
    wrote_rezumat_label = False
    wrote_body_start = False
    wrote_title_en = False
    wrote_abstract_label = False
    wrote_biblio_header = False
    # figure_numbers and table_numbers are pre-computed above

    # Collect English title and abstract text to write before bibliography
    collected_title_ro = []      # fallback: translate if no title_en found
    collected_title_en = []
    collected_abstract_text = []
    collected_rezumat_text = []  # fallback: translate if no English abstract

    # Collect consecutive items by type for proper spacing
    n = len(paragraphs)

    for p_info in paragraphs:
        idx = p_info["idx"]
        text = p_info["text"]
        sec_type = section_map.get(idx, "body")

        if sec_type == "skip":
            continue

        if sec_type == "empty":
            # Only add empty lines in body context, not between structural elements
            # We handle spacing explicitly per section transition
            continue

        # Insert authors when transitioning away from title
        if prev_type == "title_ro" and sec_type != "title_ro" and not wrote_authors:
            add_blank_lines(1)
            if authors:
                authors_line = ", ".join(authors)
                add_paragraph(authors_line, 'Author')
            else:
                # Collect all author paragraphs and join on single line
                author_texts = []
                for ap in paragraphs:
                    if section_map.get(ap["idx"]) == "author" and ap["text"].strip():
                        author_texts.append(ap["text"].strip())
                if author_texts:
                    authors_line = ", ".join(author_texts)
                    add_paragraph(authors_line, 'Author')
            wrote_authors = True

        if sec_type == "title_ro":
            if not wrote_title:
                # 6 blank lines before title (top of page padding)
                add_blank_lines(6)
                wrote_title = True
            add_paragraph(text.upper(), 'Title')
            if text.strip():
                collected_title_ro.append(text.strip())
            prev_type = sec_type
            continue

        if sec_type == "author":
            # Already handled above when transitioning from title
            prev_type = sec_type
            continue

        if sec_type == "rezumat_label":
            if not wrote_rezumat_label:
                add_blank_lines(2)
                # Always write "Rezumat" — skip combined labels like "REZUMAT / ABSTRACT"
                add_paragraph("Rezumat", 'Chapter Heading')
                add_blank_lines(1)
                wrote_rezumat_label = True
            # Skip any subsequent rezumat labels
            prev_type = sec_type
            continue

        if sec_type == "rezumat_text":
            # Italic body text for abstract
            p = add_paragraph(text, 'Normal', italic=True)
            if text.strip():
                collected_rezumat_text.append(text.strip())
            prev_type = sec_type
            continue

        if sec_type == "keywords":
            # Skip keywords paragraphs entirely
            prev_type = sec_type
            continue

        if sec_type == "heading1":
            if prev_type == "rezumat_text" and not wrote_body_start:
                wrote_body_start = True
            if prev_type is not None:
                add_blank_lines(1)
            heading_text = smart_title_case_ro(text)
            p = add_paragraph(heading_text, 'Chapter Heading')
            p.paragraph_format.keep_with_next = True
            add_blank_lines(1)
            prev_type = sec_type
            continue

        if sec_type == "heading2":
            if prev_type != "heading1":
                add_blank_lines(1)
            heading_text = sentence_case_ro(text)
            p = add_paragraph(heading_text, 'Sub Heading')
            p.paragraph_format.keep_with_next = True
            add_blank_lines(1)
            prev_type = sec_type
            continue

        if sec_type == "body":
            fixed = fix_table_refs(text)
            if p_info.get("is_bold"):
                add_paragraph(fixed, 'Normal', bold=True)
            else:
                add_paragraph(fixed, 'Normal')
            prev_type = sec_type
            continue

        if sec_type == "list_item":
            text = fix_table_refs(text)
            # Use source indent values if available, otherwise default hanging indent
            p = add_paragraph(text, 'List Paragraph')
            src_left = p_info.get("src_left_indent")
            src_first = p_info.get("src_first_indent")
            if src_left is not None:
                p.paragraph_format.left_indent = Emu(src_left)
            else:
                p.paragraph_format.left_indent = Mm(12.7)
            if src_first is not None:
                p.paragraph_format.first_line_indent = Emu(src_first)
            else:
                p.paragraph_format.first_line_indent = Mm(-6.3)
            prev_type = sec_type
            continue

        if sec_type == "formula_label":
            # Bold header of a formula box — redundant with the subchapter title, skip
            prev_type = sec_type
            continue

        if sec_type == "formula":
            formula_num = formula_numbers.get(idx)
            add_formula_paragraph(text, centered=True, formula_num=formula_num)
            prev_type = sec_type
            continue

        if sec_type == "formula_legend":
            add_formula_paragraph(text, centered=False)
            prev_type = sec_type
            continue

        if sec_type == "figure_caption":
            fig_num = figure_numbers.get(idx, 0)
            # Insert reference sentence if this figure is unreferenced
            if idx in ref_insertions:
                add_paragraph(ref_insertions[idx], 'Normal')
            # Strip existing prefix like "Figura 1.", "Fig. 2:", etc.
            caption_text = re.sub(r'^(Figura|Fig\.?)\s*\d+(?:\.\d+)?[\.:]\s*', '', text, flags=re.IGNORECASE).strip()
            add_paragraph(f"Figura {fig_num}. {caption_text}", 'Figure Caption')
            prev_type = sec_type
            continue

        if sec_type == "table_caption":
            # Skip post-captions (caption appears after its table in the source —
            # it was already written as part of the table_content handler)
            if idx in post_caption_idxs:
                prev_type = sec_type
                continue
            tbl_num = table_numbers.get(idx, 0)
            # Strip existing prefix like "Tabel 1.", "Tabel 2.2.", "Tabelul 3:", etc.
            caption_text = re.sub(r'^(Tabel|Tabelul)\s*\d+(?:\.\d+)?[\.:]\s*', '', text, flags=re.IGNORECASE).strip()
            add_paragraph(f"Tabel {tbl_num}. {caption_text}", 'Table Caption')
            prev_type = sec_type
            continue

        if sec_type == "table_content":
            # Recreate actual Word table from stored grid data
            table_data = p_info.get("table_data")
            if table_data:
                num_rows = len(table_data)
                num_cols = max(len(row) for row in table_data) if table_data else 1

                # Determine caption source: pre-caption (prev_type), post-caption, or auto-generate
                has_pre_caption = (prev_type == "table_caption")
                post_cap_idx = post_caption_map.get(idx)
                post_cap_text = None
                if post_cap_idx is not None:
                    raw = para_by_idx.get(post_cap_idx, {}).get("text", "")
                    post_cap_text = re.sub(r'^(Tabel|Tabelul)\s*\d+(?:\.\d+)?[\.:]\s*', '', raw, flags=re.IGNORECASE).strip()

                # Insert reference sentence (before the caption)
                if idx in ref_insertions and not has_pre_caption:
                    add_paragraph(ref_insertions[idx], 'Normal')

                # Add spacing before table
                add_blank_lines(1)

                if not has_pre_caption:
                    tbl_num = table_numbers.get(idx, 0)
                    if post_cap_text:
                        # Use the explicit post-caption text from the source
                        add_paragraph(f"Tabel {tbl_num}. {post_cap_text}", 'Table Caption')
                    else:
                        # Auto-generate caption for fully uncaptioned tables
                        caption_text = table_captions_map.get(idx, "")
                        if caption_text:
                            add_paragraph(f"Tabel {tbl_num}. {caption_text}", 'Table Caption')
                        else:
                            add_paragraph(f"Tabel {tbl_num}.", 'Table Caption')

                tbl = out_doc.add_table(rows=num_rows, cols=num_cols)
                tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
                tbl.style = 'Table Grid'

                for ri, row_data in enumerate(table_data):
                    for ci, cell_text in enumerate(row_data):
                        if ci < num_cols:
                            cell = tbl.rows[ri].cells[ci]
                            cell.text = cell_text
                            # Format cell paragraphs
                            for cp in cell.paragraphs:
                                cp.paragraph_format.first_line_indent = Mm(0)
                                cp.paragraph_format.space_before = Pt(0)
                                cp.paragraph_format.space_after = Pt(0)
                                cp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                for run in cp.runs:
                                    run.font.name = "Times New Roman"
                                    run.font.size = Pt(11)
                                    # Bold header row
                                    if ri == 0:
                                        run.font.bold = True

                add_blank_lines(1)
            prev_type = sec_type
            continue

        if sec_type == "title_en":
            # Collect for deferred output before bibliography
            collected_title_en.append(text)
            prev_type = sec_type
            continue

        if sec_type == "abstract_label":
            # Skip — we'll write our own label when outputting abstract
            prev_type = sec_type
            continue

        if sec_type == "abstract_text":
            # Collect for deferred output before bibliography
            collected_abstract_text.append(text)
            prev_type = sec_type
            continue

        if sec_type == "bibliography_header":
            if not wrote_biblio_header:
                # Write English title + abstract before bibliography
                en_title = title_en if title_en else (" ".join(collected_title_en) if collected_title_en else None)

                # If no English title found, translate Romanian title
                if not en_title and collected_title_ro:
                    print("No English title found — translating title...")
                    translated = translate_to_english(collected_title_ro, model=model)
                    en_title = " ".join(translated)

                # If no English abstract found, translate Rezumat
                abstract_text = collected_abstract_text
                if not abstract_text and collected_rezumat_text:
                    print("No English abstract found — translating Rezumat...")
                    abstract_text = translate_to_english(collected_rezumat_text, model=model)

                if en_title or abstract_text:
                    if en_title:
                        add_blank_lines(2)
                        add_paragraph(en_title.upper(), 'Title', bold=False)
                        add_blank_lines(2)
                    # Abstract label
                    if abstract_text:
                        add_paragraph("Abstract", 'Chapter Heading')
                        add_blank_lines(1)
                        for at in abstract_text:
                            add_paragraph(at, 'Normal', italic=True)
                        add_blank_lines(1)

                # Bibliography header — always "Bibliografie"
                add_blank_lines(3)
                add_paragraph("Bibliografie", 'Bibliography Header')
                add_blank_lines(1)
                wrote_biblio_header = True
            prev_type = sec_type
            continue

        if sec_type == "bibliography_entry":
            bib_pos = bib_idx_map.get(idx)
            if bib_pos is not None and bib_pos < len(bib_parsed):
                # IEEE-style: multiple runs with different formatting
                p = out_doc.add_paragraph(style='Bibliography Entry')
                for part in bib_parsed[bib_pos]:
                    run = p.add_run(part["text"])
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)
                    if part["format"] == "bold":
                        run.font.bold = True
                    elif part["format"] == "italic":
                        run.font.italic = True
            else:
                add_paragraph(text, 'Bibliography Entry', bold=True)
            prev_type = sec_type
            continue

        # Fallback
        add_paragraph(text, 'Normal')
        prev_type = sec_type

    return out_doc


def main():
    parser = argparse.ArgumentParser(
        description="Format a .docx document for conference submission using Claude AI"
    )
    parser.add_argument("input", help="Input .docx file path")
    parser.add_argument("-o", "--output", help="Output file path (default: input_formatted.docx)")
    parser.add_argument("--model", default="claude-sonnet-4-6",
                        help="Claude model to use (default: claude-sonnet-4-6)")
    parser.add_argument("--authors", nargs="+",
                        help="Author names (e.g. --authors 'Gheorghe Badea' 'George Naghiu')")
    parser.add_argument("--title-en",
                        help="English translation of the title (added before bibliography)")
    parser.add_argument("--dry-run", action="store_true",
                        help="Only classify paragraphs, don't generate output")
    parser.add_argument("--show-classification", action="store_true",
                        help="Print the full classification before generating output")
    parser.add_argument("--skip-diacritics", action="store_true",
                        help="Skip diacritics restoration (document already has correct Romanian chars)")
    parser.add_argument("--load-classification", metavar="PATH",
                        help="Load a saved classification JSON, skipping Claude classification")

    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: File not found: {input_path}")
        sys.exit(1)

    if args.output:
        output_path = Path(args.output)
    else:
        output_path = input_path.with_stem(input_path.stem + "_formatted")

    template_path = str(Path(__file__).parent / "template_conference.docx")

    # Step 1: Extract paragraphs
    print(f"Reading: {input_path}")
    paragraphs, doc = extract_paragraphs(str(input_path))
    print(f"Extracted {len(paragraphs)} paragraphs")

    # Step 2: Restore Romanian diacritics (skip if doc is already clean)
    if args.skip_diacritics:
        print("Skipping diacritics restoration (--skip-diacritics)")
    else:
        paragraphs = restore_diacritics(paragraphs, model=args.model)

    # Step 3: Classify with Claude (or load cached classification)
    if args.load_classification:
        print(f"Loading classification from: {args.load_classification}")
        with open(args.load_classification, encoding="utf-8") as f:
            raw = json.load(f)
        section_map = {int(k): v for k, v in raw.items()}
        # Fill any indices present in paragraphs but missing from cache
        for p in paragraphs:
            if p["idx"] not in section_map:
                section_map[p["idx"]] = "empty" if p["is_empty"] else "body"
        print(f"Loaded {len(section_map)} classifications")
    else:
        section_map = classify_with_claude(paragraphs, model=args.model, docx_path=str(input_path))
        # Auto-save sidecar for future reuse
        classification_cache = output_path.with_suffix(".classification.json")
        with open(classification_cache, "w", encoding="utf-8") as f:
            json.dump(section_map, f)
        print(f"Classification cached: {classification_cache}")

    # Show classification if requested
    if args.show_classification or args.dry_run:
        print("\n--- Classification ---")
        for p in paragraphs:
            sec = section_map.get(p["idx"], "???")
            text = p["text"][:70] if p["text"] else "[EMPTY]"
            print(f"  P{p['idx']:3d} [{sec:22s}] {text}")
        print("--- End Classification ---\n")

    if args.dry_run:
        print("Dry run complete. No output file generated.")
        return

    # Step 3: Build formatted document
    print(f"\nBuilding formatted document...")
    out_doc = build_formatted_document(str(input_path), section_map, paragraphs, template_path,
                                       model=args.model, authors=args.authors, title_en=args.title_en,
                                       fast=args.skip_diacritics)

    # Step 4: Save
    out_doc.save(str(output_path))
    print(f"\nDone! Output saved to: {output_path}")
    print(f"Total paragraphs in output: {len(out_doc.paragraphs)}")


if __name__ == "__main__":
    main()
