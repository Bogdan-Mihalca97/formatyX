"""
Creates the conference formatting template .docx with all styles pre-configured.
This template is used by the formatter to produce correctly styled output.

Based on the reference document: V3_ANALIZA ECONOMICA...
"""

from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


FONT_NAME = "Times New Roman"

# Page: A5
PAGE_WIDTH = Mm(148)
PAGE_HEIGHT = Mm(210)
MARGIN = Mm(20)
HEADER_FOOTER_DIST = Mm(12.7)
TAB_INDENT = Mm(12.7)


def set_doc_defaults(doc):
    """Set document-level default font to Times New Roman via XML."""
    styles_element = doc.styles.element
    doc_defaults = styles_element.find(qn('w:docDefaults'))
    if doc_defaults is None:
        doc_defaults = parse_xml(
            f'<w:docDefaults {nsdecls("w")}>'
            f'  <w:rPrDefault>'
            f'    <w:rPr>'
            f'      <w:rFonts w:ascii="{FONT_NAME}" w:eastAsia="{FONT_NAME}" w:hAnsi="{FONT_NAME}" w:cs="{FONT_NAME}"/>'
            f'      <w:lang w:val="ro-RO" w:eastAsia="en-US" w:bidi="ar-SA"/>'
            f'    </w:rPr>'
            f'  </w:rPrDefault>'
            f'  <w:pPrDefault/>'
            f'</w:docDefaults>'
        )
        styles_element.insert(0, doc_defaults)
    else:
        rpr_default = doc_defaults.find(qn('w:rPrDefault'))
        if rpr_default is None:
            rpr_default = parse_xml(
                f'<w:rPrDefault {nsdecls("w")}>'
                f'  <w:rPr>'
                f'    <w:rFonts w:ascii="{FONT_NAME}" w:eastAsia="{FONT_NAME}" w:hAnsi="{FONT_NAME}" w:cs="{FONT_NAME}"/>'
                f'  </w:rPr>'
                f'</w:rPrDefault>'
            )
            doc_defaults.insert(0, rpr_default)
        else:
            rpr = rpr_default.find(qn('w:rPr'))
            if rpr is None:
                rpr = parse_xml(
                    f'<w:rPr {nsdecls("w")}>'
                    f'  <w:rFonts w:ascii="{FONT_NAME}" w:eastAsia="{FONT_NAME}" w:hAnsi="{FONT_NAME}" w:cs="{FONT_NAME}"/>'
                    f'</w:rPr>'
                )
                rpr_default.append(rpr)
            else:
                rfonts = rpr.find(qn('w:rFonts'))
                if rfonts is None:
                    rfonts = parse_xml(
                        f'<w:rFonts {nsdecls("w")} w:ascii="{FONT_NAME}" w:eastAsia="{FONT_NAME}" w:hAnsi="{FONT_NAME}" w:cs="{FONT_NAME}"/>'
                    )
                    rpr.insert(0, rfonts)
                else:
                    rfonts.set(qn('w:ascii'), FONT_NAME)
                    rfonts.set(qn('w:eastAsia'), FONT_NAME)
                    rfonts.set(qn('w:hAnsi'), FONT_NAME)
                    rfonts.set(qn('w:cs'), FONT_NAME)


def configure_style(style, font_size=None, bold=None, italic=None,
                    alignment=None, first_line_indent=None, left_indent=None,
                    space_before=None, space_after=None, line_spacing=None,
                    font_name=None, color=None):
    """Configure a style's font and paragraph properties."""
    f = style.font
    if font_name:
        f.name = font_name
        # Clear theme font references so the explicit font name takes effect
        rpr = style.element.find(qn('w:rPr'))
        if rpr is not None:
            rfonts = rpr.find(qn('w:rFonts'))
            if rfonts is not None:
                for attr in ['w:asciiTheme', 'w:hAnsiTheme', 'w:eastAsiaTheme', 'w:cstheme']:
                    rfonts.attrib.pop(qn(attr), None)
    if font_size is not None:
        f.size = font_size
    if bold is not None:
        f.bold = bold
    if italic is not None:
        f.italic = italic
    if color is not None:
        f.color.rgb = color

    pf = style.paragraph_format
    if alignment is not None:
        pf.alignment = alignment
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    if left_indent is not None:
        pf.left_indent = left_indent
    if space_before is not None:
        pf.space_before = space_before
    if space_after is not None:
        pf.space_after = space_after
    if line_spacing is not None:
        pf.line_spacing = line_spacing
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE


def create_template(output_path="template_conference.docx"):
    doc = Document()

    # --- Page setup ---
    for section in doc.sections:
        section.page_width = PAGE_WIDTH
        section.page_height = PAGE_HEIGHT
        section.top_margin = MARGIN
        section.bottom_margin = MARGIN
        section.left_margin = MARGIN
        section.right_margin = MARGIN
        section.header_distance = HEADER_FOOTER_DIST
        section.footer_distance = HEADER_FOOTER_DIST
        section.orientation = WD_ORIENT.PORTRAIT

    # --- Document defaults (Times New Roman at XML level) ---
    set_doc_defaults(doc)

    # --- Normal style ---
    # Body text: 11pt, justify, first-line indent 12.7mm, single spacing
    normal = doc.styles['Normal']
    configure_style(normal,
                    font_name=FONT_NAME,
                    font_size=Pt(11),
                    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    first_line_indent=TAB_INDENT,
                    space_before=Pt(0),
                    space_after=Pt(0),
                    line_spacing=1.0)

    # --- Title style ---
    # 12pt, bold, center, no indent
    title_style = doc.styles['Title']
    configure_style(title_style,
                    font_name=FONT_NAME,
                    font_size=Pt(12),
                    bold=True,
                    italic=False,
                    alignment=WD_ALIGN_PARAGRAPH.CENTER,
                    first_line_indent=Mm(0),
                    left_indent=Mm(0),
                    space_before=Pt(0),
                    space_after=Pt(0),
                    line_spacing=1.0,
                    color=RGBColor(0, 0, 0))
    # Remove any underline/border from Title
    title_style.font.underline = False
    # Remove the bottom border (blue line) that Word's built-in Title style adds
    pPr = title_style.element.find(qn('w:pPr'))
    if pPr is not None:
        pBdr = pPr.find(qn('w:pBdr'))
        if pBdr is not None:
            pPr.remove(pBdr)
    # Also strip it via explicit empty border to prevent theme inheritance
    pPr = title_style.element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    # --- Chapter Heading style (custom, replaces built-in Heading 1) ---
    # 11pt, bold, left-aligned, hanging indent (number at margin, text at TAB)
    try:
        h1 = doc.styles['Chapter Heading']
    except KeyError:
        h1 = doc.styles.add_style('Chapter Heading', WD_STYLE_TYPE.PARAGRAPH)
    h1.base_style = doc.styles['Normal']
    configure_style(h1,
                    font_name=FONT_NAME,
                    font_size=Pt(11),
                    bold=True,
                    italic=False,
                    alignment=WD_ALIGN_PARAGRAPH.LEFT,
                    left_indent=TAB_INDENT,
                    first_line_indent=Mm(0),
                    space_before=Pt(0),
                    space_after=Pt(0),
                    line_spacing=1.0,
                    color=RGBColor(0, 0, 0))

    # --- Sub Heading style (custom, replaces built-in Heading 2) ---
    # Bold, same indent pattern as Chapter Heading
    try:
        h2 = doc.styles['Sub Heading']
    except KeyError:
        h2 = doc.styles.add_style('Sub Heading', WD_STYLE_TYPE.PARAGRAPH)
    h2.base_style = doc.styles['Normal']
    configure_style(h2,
                    font_name=FONT_NAME,
                    font_size=Pt(11),
                    bold=True,
                    italic=False,
                    alignment=WD_ALIGN_PARAGRAPH.LEFT,
                    left_indent=TAB_INDENT,
                    first_line_indent=Mm(0),
                    space_before=Pt(0),
                    space_after=Pt(0),
                    line_spacing=1.0,
                    color=RGBColor(0, 0, 0))

    # --- List Paragraph style ---
    # 11pt, justify, no first-line indent by default
    try:
        list_para = doc.styles['List Paragraph']
    except KeyError:
        list_para = doc.styles.add_style('List Paragraph', WD_STYLE_TYPE.PARAGRAPH)
    configure_style(list_para,
                    font_name=FONT_NAME,
                    font_size=Pt(11),
                    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    first_line_indent=Mm(0),
                    left_indent=Mm(0),
                    space_before=Pt(0),
                    space_after=Pt(0),
                    line_spacing=1.0)

    # --- Custom styles for conference-specific elements ---

    # Author style: 11pt, normal, center, no indent
    try:
        author_style = doc.styles['Author']
    except KeyError:
        author_style = doc.styles.add_style('Author', WD_STYLE_TYPE.PARAGRAPH)
    author_style.base_style = doc.styles['Normal']
    configure_style(author_style,
                    font_name=FONT_NAME,
                    font_size=Pt(11),
                    bold=False,
                    alignment=WD_ALIGN_PARAGRAPH.CENTER,
                    first_line_indent=Mm(0),
                    left_indent=Mm(0))

    # Rezumat/Abstract label: bold, 11pt, justify, indent 1 TAB
    # (uses Heading 1 style in the reference doc)

    # Abstract text: italic, 11pt, justify, first-line indent
    try:
        abstract_style = doc.styles['Abstract Text']
    except KeyError:
        abstract_style = doc.styles.add_style('Abstract Text', WD_STYLE_TYPE.PARAGRAPH)
    abstract_style.base_style = doc.styles['Normal']
    configure_style(abstract_style,
                    font_name=FONT_NAME,
                    font_size=Pt(11),
                    italic=True,
                    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    first_line_indent=TAB_INDENT)

    # Figure caption: 9pt, center, no indent
    try:
        fig_caption = doc.styles['Figure Caption']
    except KeyError:
        fig_caption = doc.styles.add_style('Figure Caption', WD_STYLE_TYPE.PARAGRAPH)
    fig_caption.base_style = doc.styles['Normal']
    configure_style(fig_caption,
                    font_name=FONT_NAME,
                    font_size=Pt(9),
                    alignment=WD_ALIGN_PARAGRAPH.CENTER,
                    first_line_indent=Mm(0),
                    left_indent=Mm(0))

    # Table caption: 10pt, bold, center, no indent
    try:
        tbl_caption = doc.styles['Table Caption']
    except KeyError:
        tbl_caption = doc.styles.add_style('Table Caption', WD_STYLE_TYPE.PARAGRAPH)
    tbl_caption.base_style = doc.styles['Normal']
    configure_style(tbl_caption,
                    font_name=FONT_NAME,
                    font_size=Pt(10),
                    bold=True,
                    alignment=WD_ALIGN_PARAGRAPH.CENTER,
                    first_line_indent=Mm(0),
                    left_indent=Mm(0))

    # Bibliography header: 12pt, bold, center
    try:
        bib_header = doc.styles['Bibliography Header']
    except KeyError:
        bib_header = doc.styles.add_style('Bibliography Header', WD_STYLE_TYPE.PARAGRAPH)
    bib_header.base_style = doc.styles['Normal']
    configure_style(bib_header,
                    font_name=FONT_NAME,
                    font_size=Pt(12),
                    bold=True,
                    alignment=WD_ALIGN_PARAGRAPH.CENTER,
                    first_line_indent=Mm(0),
                    left_indent=Mm(0))

    # Bibliography entry: 11pt, justify, no first-line indent
    try:
        bib_entry = doc.styles['Bibliography Entry']
    except KeyError:
        bib_entry = doc.styles.add_style('Bibliography Entry', WD_STYLE_TYPE.PARAGRAPH)
    bib_entry.base_style = doc.styles['Normal']
    configure_style(bib_entry,
                    font_name=FONT_NAME,
                    font_size=Pt(11),
                    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    first_line_indent=Mm(0),
                    left_indent=Mm(0))

    # No Spacing style (for blank lines)
    try:
        no_spacing = doc.styles['No Spacing']
    except KeyError:
        no_spacing = doc.styles.add_style('No Spacing', WD_STYLE_TYPE.PARAGRAPH)
    configure_style(no_spacing,
                    font_name=FONT_NAME,
                    font_size=Pt(11),
                    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    first_line_indent=TAB_INDENT,
                    space_before=Pt(0),
                    space_after=Pt(0),
                    line_spacing=1.0)

    # Remove the default empty paragraph
    if doc.paragraphs:
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)

    doc.save(output_path)
    print(f"Template saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    create_template()
