"""
AI-powered academic paper generator using Claude.
Generates Romanian conference papers section by section with context awareness.
"""

import os
import requests
from anthropic import Anthropic
from dotenv import load_dotenv
import re
from docx import Document
from docx.shared import Mm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

load_dotenv()

# Sections marked optional=True will be skipped if Claude returns "N/A"
SECTIONS = [
    {
        "key": "title",
        "label": "Titlu",
        "optional": False,
        "max_tokens": 120,
        "instruction": (
            "Generează un titlu academic clar, precis și concis în limba română pentru această lucrare de cercetare. "
            "Titlul trebuie să reflecte exact subiectul și să fie potrivit pentru o conferință științifică. "
            "Returnează DOAR titlul, fără prefixe, ghilimele sau formatare suplimentară."
        ),
    },
    {
        "key": "rezumat",
        "label": "Rezumat",
        "optional": False,
        "max_tokens": 700,
        "instruction": (
            "Scrie un rezumat academic de aproximativ 200-250 de cuvinte în limba română. "
            "Prezintă: contextul cercetării, obiectivele, metodologia utilizată, principalele rezultate și concluzii. "
            "Scrie în stil academic formal. Nu include anteturi, prefixe sau text introductiv. "
            "Returnează DOAR textul rezumatului."
        ),
    },
    {
        "key": "keywords_ro",
        "label": "Cuvinte cheie",
        "optional": False,
        "max_tokens": 120,
        "instruction": (
            "Generează 5-7 cuvinte cheie relevante în română pentru această lucrare. "
            "Returnează DOAR cuvintele cheie separate prin virgulă și spațiu, fără alte texte."
        ),
    },
    {
        "key": "nomenclature",
        "label": "Nomenclatură",
        "optional": True,
        "max_tokens": 1800,
        "instruction": (
            "Dacă lucrarea implică simboluri matematice, mărimi fizice sau abrevieri tehnice specifice, "
            "listează-le câte unul pe linie în formatul: Simbol — Definiție [unitate]. "
            "IMPORTANT: folosește notația cu underscore pentru indici (ex: Ex_in, eta_ex, ex_ph, Q_rec, m_comb) "
            "astfel încât indicii să poată fi formatați ca subscript în document. "
            "Dacă lucrarea nu necesită nomenclatură, returnează exact 'N/A'. "
            "Nu include texte introductive sau anteturi. Listează TOATE simbolurile relevante."
        ),
    },
    {
        "key": "introduction",
        "label": "1. Introducere",
        "optional": False,
        "max_tokens": 1800,
        "instruction": (
            "Scrie introducerea lucrării în română (450-550 cuvinte). "
            "Include: contextul teoretic și starea artei, motivația cercetării, "
            "relevanța temei în domeniu, obiectivele specifice și structura lucrării. "
            "Folosește stil academic formal. Nu include anteturi. Returnează DOAR conținutul introducerii."
        ),
    },
    {
        "key": "relevance",
        "label": "2. Relevanța cercetării",
        "optional": False,
        "max_tokens": 1200,
        "instruction": (
            "Descrie relevanța practică și teoretică a cercetării în română (300-400 cuvinte). "
            "Include: importanța domeniului, impactul potențial al rezultatelor, "
            "contribuțiile originale față de literatura existentă. "
            "Nu include anteturi. Returnează DOAR conținutul secțiunii."
        ),
    },
    {
        "key": "methodology",
        "label": "3. Metodologie (Model Matematic)",
        "optional": False,
        "max_tokens": 4000,
        "instruction": (
            "Scrie metodologia completă în română cu subcapitole, ecuații și tabele. "
            "Structurează cu subcapitole folosind formatul '### 3.x Titlu subcapitol'. "
            "Include 3-4 subcapitole (ex: descrierea sistemului, modelul matematic, ipoteze, parametri). "
            "Pentru ecuații matematice, scrie fiecare ecuație pe o linie separată folosind '=' și notație cu underscore pentru indici (ex: Ex_in = m_dot * ex_ph). "
            "Include cel puțin un tabel de parametri tehnici în format markdown: "
            "| Parametru | Simbol | Valoare | Unitate | pe prima linie, "
            "|---|---|---|---| pe a doua linie, apoi rândurile de date. "
            "Nu include 'Metodologie' ca antet principal — începe direct cu primul subcapitol '### 3.1 ...'."
        ),
    },
    {
        "key": "materials_methods",
        "label": "4. Materiale și Metode",
        "optional": False,
        "max_tokens": 3000,
        "instruction": (
            "Descrie materialele și metodele în română cu subcapitole și tabele unde este relevant. "
            "Folosește formatul '### 4.x Titlu' pentru subcapitole (2-3 subcapitole). "
            "Dacă există echipamente sau materiale cu specificații tehnice, include un tabel markdown: "
            "| Echipament/Material | Specificații | Producător/Standard | "
            "|---|---|---| urmat de rândurile de date. "
            "Folosește voce pasivă și limbaj tehnic precis. "
            "Nu include 'Materiale și Metode' ca antet — începe direct cu '### 4.1 ...'."
        ),
    },
    {
        "key": "technology_overview",
        "label": "5. Prezentare Tehnologică",
        "optional": True,
        "max_tokens": 1600,
        "instruction": (
            "Dacă lucrarea implică o tehnologie specifică, prezintă o privire de ansamblu asupra "
            "tehnologiilor relevante în română (350-450 cuvinte): principii de funcționare, "
            "avantaje/dezavantaje comparativ cu alternativele, stadiul actual al tehnologiei. "
            "Dacă subiectul nu implică o componentă tehnologică distinctă față de metodologie, "
            "returnează exact 'N/A'. Nu include anteturi."
        ),
    },
    {
        "key": "case_study",
        "label": "6. Studiu de Caz",
        "optional": True,
        "max_tokens": 1800,
        "instruction": (
            "Dacă lucrarea include un studiu de caz specific sau o aplicație concretă, "
            "descrie-l în română (450-550 cuvinte): contextul specific, datele de intrare, "
            "rezultatele aplicației practice, analiza performanței. "
            "Dacă subiectul nu se pretează unui studiu de caz distinct față de rezultate, "
            "returnează exact 'N/A'. Nu include anteturi."
        ),
    },
    {
        "key": "results",
        "label": "7. Rezultate și Discuții",
        "optional": False,
        "max_tokens": 4000,
        "instruction": (
            "Prezintă și discută rezultatele în română cu subcapitole și tabele de rezultate. "
            "Folosește formatul '### 7.x Titlu' pentru subcapitole (2-3 subcapitole). "
            "Include cel puțin un tabel de rezultate în format markdown cu valorile calculate/măsurate: "
            "| Mărime | Simbol | Valoare calculată | Unitate | pe prima linie, "
            "|---|---|---|---| pe a doua, apoi datele. "
            "Discută semnificația rezultatelor și compară cu literatura. "
            "Nu include 'Rezultate' ca antet — începe direct cu '### 7.1 ...'."
        ),
    },
    {
        "key": "standards",
        "label": "8. Standarde și Reglementări",
        "optional": True,
        "max_tokens": 1200,
        "instruction": (
            "Dacă lucrarea este relevantă pentru standarde tehnice, reglementări sau norme specifice, "
            "discută-le în română (300-400 cuvinte): standardele aplicabile, cerințele de conformitate, "
            "implicațiile pentru proiectare sau implementare. "
            "Dacă nu există standarde specific aplicabile subiectului, returnează exact 'N/A'. "
            "Nu include anteturi."
        ),
    },
    {
        "key": "future_challenges",
        "label": "9. Provocări Viitoare",
        "optional": False,
        "max_tokens": 1200,
        "instruction": (
            "Identifică și discută provocările viitoare și direcțiile de cercetare în română (300-400 cuvinte). "
            "Include: limitările actuale ale cercetării, oportunități de îmbunătățire, "
            "direcții viitoare de cercetare și dezvoltare. "
            "Nu include anteturi. Returnează DOAR conținutul secțiunii."
        ),
    },
    {
        "key": "environmental",
        "label": "10. Sustenabilitate și Impact de Mediu",
        "optional": True,
        "max_tokens": 1200,
        "instruction": (
            "Dacă lucrarea are implicații de mediu sau sustenabilitate relevante, "
            "analizează-le în română (300-400 cuvinte): impactul de mediu al soluțiilor propuse, "
            "beneficiile de sustenabilitate, amprenta de carbon sau eficiența energetică. "
            "Dacă subiectul nu are o componentă de mediu semnificativă, returnează exact 'N/A'. "
            "Nu include anteturi."
        ),
    },
    {
        "key": "conclusions",
        "label": "11. Concluzii",
        "optional": False,
        "max_tokens": 1000,
        "instruction": (
            "Scrie concluziile lucrării în română (200-280 cuvinte). "
            "Sintetizează: principalele descoperiri, contribuțiile originale, "
            "limitările cercetării și recomandările pentru viitor. "
            "Nu include anteturi. Returnează DOAR textul concluziilor."
        ),
    },
    {
        "key": "title_en",
        "label": "Title (EN)",
        "optional": False,
        "max_tokens": 120,
        "instruction": (
            "Translate the Romanian paper title into English. "
            "Keep it precise and academic. Return ONLY the translated title, no quotes or extra text."
        ),
    },
    {
        "key": "abstract_en",
        "label": "Abstract (EN)",
        "optional": False,
        "max_tokens": 700,
        "instruction": (
            "Translate the Romanian abstract (Rezumat) into English (200-250 words). "
            "Maintain academic tone and precision. Return ONLY the translated abstract text."
        ),
    },
    {
        "key": "keywords_en",
        "label": "Keywords (EN)",
        "optional": False,
        "max_tokens": 120,
        "instruction": (
            "Translate the Romanian keywords into English. "
            "Return ONLY the translated keywords separated by commas, no other text."
        ),
    },
    {
        "key": "references",
        "label": "Bibliografie",
        "optional": False,
        "max_tokens": 1800,
        "instruction": (
            "Generează 10-14 referințe bibliografice relevante și realiste în format IEEE pentru această lucrare. "
            "Referințele trebuie să fie din domeniul temei cercetate, publicate în reviste sau conferințe recunoscute. "
            "Returnează DOAR lista numerotată de referințe, fără alte texte."
        ),
    },
]

SECTION_MAP = {s["key"]: s for s in SECTIONS}
SECTION_KEYS = [s["key"] for s in SECTIONS]

SYSTEM_PROMPT = (
    "Ești un expert în redactarea lucrărilor academice românești pentru conferințe științifice de inginerie și tehnologie. "
    "Scrii în stil academic formal, precis și coerent. Menții consistența între secțiuni. "
    "Folosești terminologie tehnică corectă și referințe la standarde internaționale când este cazul. "
    "Nu adaugi niciodată anteturi, prefixe sau comentarii — returnezi DOAR conținutul cerut. "
    "Dacă o secțiune opțională nu este aplicabilă, returnezi exact 'N/A'."
)


def _build_context(generated: dict) -> str:
    """Build a context string from already-generated sections (excluding N/A ones)."""
    parts = []
    for key in SECTION_KEYS:
        text = generated.get(key, "")
        if text and text.strip() != "N/A":
            label = SECTION_MAP[key]["label"]
            parts.append(f"[{label}]\n{text}")
    return "\n\n".join(parts)


def generate_section(key: str, topic: str, domain: str, objectives: str,
                     keywords: str, generated: dict,
                     model: str = "claude-sonnet-4-6") -> str:
    """Generate a single section using Claude with full context."""
    client = Anthropic()
    section = SECTION_MAP[key]

    context = _build_context(generated)
    context_block = f"\n\nSecțiuni deja generate:\n{context}" if context else ""

    user_msg = (
        f"Subiect lucrare: {topic}\n"
        f"Domeniu: {domain}\n"
        f"Obiective cercetare: {objectives or 'nespecificate'}\n"
        f"Cuvinte cheie sugerate: {keywords or 'nespecificate'}"
        f"{context_block}\n\n"
        f"Sarcină: {section['instruction']}"
    )

    response = client.messages.create(
        model=model,
        max_tokens=section["max_tokens"],
        system=SYSTEM_PROMPT,
        messages=[{"role": "user", "content": user_msg}],
    )
    text = response.content[0].text.strip()

    # Treat near-empty or explicit N/A as not applicable
    if section.get("optional") and text.upper() in ("N/A", "NA", "N/A.", "-", ""):
        return "N/A"
    return text


def check_grammar(text: str, language: str = "ro-RO") -> list[dict]:
    """Check grammar using LanguageTool free API (20k char limit)."""
    MAX_CHARS = 19000
    text = text[:MAX_CHARS]
    try:
        resp = requests.post(
            "https://api.languagetool.org/v2/check",
            data={"text": text, "language": language, "enabledOnly": "false"},
            headers={"Content-Type": "application/x-www-form-urlencoded"},
            timeout=20,
        )
        resp.raise_for_status()
        matches = resp.json().get("matches", [])
        return [
            {
                "message": m["message"],
                "offset": m["offset"],
                "length": m["length"],
                "replacements": [r["value"] for r in m["replacements"][:3]],
                "context": m["context"]["text"],
            }
            for m in matches
        ]
    except requests.HTTPError as e:
        return [{"message": f"Grammar check error ({e.response.status_code}): {e.response.text[:200]}", "offset": 0, "length": 0, "replacements": [], "context": ""}]
    except Exception as e:
        return [{"message": f"Grammar check unavailable: {e}", "offset": 0, "length": 0, "replacements": [], "context": ""}]


def build_docx(paper: dict, output_path: str):
    """Build a conference-formatted DOCX matching the formatter's exact rules."""
    from create_template import create_template

    if not os.path.exists("template_conference.docx"):
        create_template("template_conference.docx")

    doc = Document("template_conference.docx")
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    _subscript_re = re.compile(r'([^\s_]+)_(\{[^}]+\}|\w+)')
    _formula_re   = re.compile(r'[=+\-*/^∑∫∂·×]')
    _FONT = "Times New Roman"

    def blank(n=1):
        for _ in range(n):
            p = doc.add_paragraph("", style='Normal')
            p.paragraph_format.first_line_indent = Mm(0)

    def _add_subscript_runs(p, text, bold=False, italic=False):
        """Add runs with subscript formatting for base_sub patterns."""
        last = 0
        for m in _subscript_re.finditer(text):
            if m.start() > last:
                r = p.add_run(text[last:m.start()])
                r.font.name = _FONT; r.font.size = Pt(11)
                r.font.bold = bold; r.font.italic = italic
            # base part
            r = p.add_run(m.group(1))
            r.font.name = _FONT; r.font.size = Pt(11)
            r.font.bold = bold; r.font.italic = italic
            # subscript part (strip braces if present)
            sub_text = m.group(2).strip('{}')
            r = p.add_run(sub_text)
            r.font.name = _FONT; r.font.size = Pt(9)
            r.font.subscript = True
            r.font.bold = bold; r.font.italic = italic
            last = m.end()
        if last < len(text):
            r = p.add_run(text[last:])
            r.font.name = _FONT; r.font.size = Pt(11)
            r.font.bold = bold; r.font.italic = italic

    def para(text, style, bold=False, italic=False, keep_next=False):
        p = doc.add_paragraph(style=style)
        _add_subscript_runs(p, text, bold=bold, italic=italic)
        if keep_next:
            p.paragraph_format.keep_with_next = True
        return p

    _md_sep_re  = re.compile(r'^\|[-| :]+\|$')
    _sub_heading_re = re.compile(r'^#{2,4}\s+(.+)')

    def _is_formula_line(line: str) -> bool:
        return (len(line) < 250
                and '=' in line
                and len(_formula_re.findall(line)) >= 2)

    def _render_md_table(table_lines: list):
        """Render a list of markdown pipe-table lines as a Word table."""
        rows = []
        for line in table_lines:
            if _md_sep_re.match(line):
                continue
            parts = line.split('|')
            cells = [c.strip() for c in parts[1:-1]] if len(parts) > 2 else [c.strip() for c in parts if c.strip()]
            if cells:
                rows.append(cells)
        if not rows:
            return
        cols = max(len(r) for r in rows)
        tbl = doc.add_table(rows=len(rows), cols=cols)
        tbl.style = 'Table Grid'
        for r_i, row_data in enumerate(rows):
            for c_i in range(cols):
                cell_text = row_data[c_i] if c_i < len(row_data) else ""
                cell = tbl.rows[r_i].cells[c_i]
                p = cell.paragraphs[0]
                p.clear()
                run = p.add_run(cell_text)
                run.font.name = _FONT
                run.font.size = Pt(10)
                if r_i == 0:
                    run.font.bold = True
        blank(1)

    def render_section(text: str):
        """Render a section body: subchapters, tables, formulas, normal text."""
        lines = text.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i].strip()

            # Subchapter heading (## or ###)
            m = _sub_heading_re.match(line)
            if m:
                blank(1)
                p = doc.add_paragraph(style='Sub Heading')
                _add_subscript_runs(p, m.group(1).strip())
                p.paragraph_format.keep_with_next = True
                blank(1)
                i += 1
                continue

            # Markdown table block — collect all consecutive pipe lines
            if line.startswith('|'):
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith('|'):
                    table_lines.append(lines[i].strip())
                    i += 1
                _render_md_table(table_lines)
                continue

            # Empty line
            if not line:
                i += 1
                continue

            # Formula line
            if _is_formula_line(line):
                p = doc.add_paragraph(style='Normal')
                p.paragraph_format.first_line_indent = Mm(0)
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _add_subscript_runs(p, line)
                i += 1
                continue

            # Normal body text
            para(line, 'Normal')
            i += 1

    def is_na(key):
        return (paper.get(key) or "").strip().upper() in ("N/A", "NA", "")

    # ── Title (RO) — 6 blank lines before, UPPERCASE ──
    if not is_na("title"):
        blank(6)
        para(paper["title"].upper(), 'Title')

    # ── Authors — 1 blank line after title, all on one line ──
    authors = paper.get("authors") or []
    if isinstance(authors, str):
        authors = [a.strip() for a in authors.splitlines() if a.strip()]
    blank(1)
    if authors:
        para(", ".join(authors), 'Author')

    # ── Rezumat — 2 blank lines before label, 1 blank after label ──
    if not is_na("rezumat"):
        blank(2)
        para("Rezumat", 'Chapter Heading', keep_next=True)
        blank(1)
        para(paper["rezumat"], 'Normal', italic=True)

    # ── Cuvinte cheie (RO) ──
    if not is_na("keywords_ro"):
        p = doc.add_paragraph(style='Normal')
        p.paragraph_format.first_line_indent = Mm(12.7)
        r = p.add_run("Cuvinte cheie: ")
        r.font.bold = True
        r.font.italic = True
        kwr = p.add_run(paper["keywords_ro"])
        kwr.font.italic = True

    # ── Nomenclatură (optional) — subscript formatting per entry ──
    if not is_na("nomenclature"):
        blank(1)
        para("Nomenclatură", 'Chapter Heading', keep_next=True)
        blank(1)
        for line in paper["nomenclature"].split("\n"):
            line = line.strip()
            if not line:
                continue
            p = doc.add_paragraph(style='Normal')
            p.paragraph_format.first_line_indent = Mm(0)
            _add_subscript_runs(p, line)

    # ── Body sections — 1 blank before heading, heading with keep_next, 1 blank after ──
    BODY_SECTIONS = [
        ("introduction",        "1. Introducere"),
        ("relevance",           "2. Relevanța Cercetării"),
        ("methodology",         "3. Metodologie (Model Matematic)"),
        ("materials_methods",   "4. Materiale și Metode"),
        ("technology_overview", "5. Prezentare Tehnologică"),
        ("case_study",          "6. Studiu de Caz"),
        ("results",             "7. Rezultate și Discuții"),
        ("standards",           "8. Standarde și Reglementări"),
        ("future_challenges",   "9. Provocări Viitoare"),
        ("environmental",       "10. Sustenabilitate și Impact de Mediu"),
        ("conclusions",         "11. Concluzii"),
    ]

    for key, heading in BODY_SECTIONS:
        if not is_na(key):
            blank(1)
            para(heading, 'Chapter Heading', keep_next=True)
            blank(1)
            render_section(paper[key])

    # ── English title — 2 blank before, UPPERCASE, 2 blank after ──
    if not is_na("title_en"):
        blank(2)
        para(paper["title_en"].upper(), 'Title')
        blank(2)

    # ── Abstract (EN) — label + 1 blank + italic text + 1 blank ──
    if not is_na("abstract_en"):
        para("Abstract", 'Chapter Heading', keep_next=True)
        blank(1)
        para(paper["abstract_en"], 'Normal', italic=True)
        blank(1)

    # ── Keywords (EN) ──
    if not is_na("keywords_en"):
        p = doc.add_paragraph(style='Normal')
        p.paragraph_format.first_line_indent = Mm(12.7)
        r = p.add_run("Keywords: ")
        r.font.bold = True
        r.font.italic = True
        kwr = p.add_run(paper["keywords_en"])
        kwr.font.italic = True

    # ── Bibliography — 3 blank before header, 1 blank after ──
    if not is_na("references"):
        blank(3)
        para("Bibliografie", 'Bibliography Header')
        blank(1)
        for line in paper["references"].split("\n"):
            line = line.strip()
            if line:
                para(line, 'Bibliography Entry')

    doc.save(output_path)
