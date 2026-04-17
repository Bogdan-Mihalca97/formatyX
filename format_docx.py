"""
Conference Document Formatter
Formats .docx files according to strict conference formatting rules.

Rules:
- Page: ISO B5 (176x250.1mm), margins 20mm all sides, header/footer 12.7mm
- Font: Times New Roman, Size 11, single spacing throughout
- TAB indent: 12.7mm (first-line indent)
- Structure order:
    1. Title (bold, 12pt, UPPERCASE, center) + 6 blank lines
    2. Authors (normal, 11pt, center) + 1 blank line
    3. "Rezumat" label (bold, 11pt, justify, indent 1 TAB)
    4. Romanian abstract (italic, 11pt, justify, indent 1 TAB) + 2 blank lines
    5. Body text (normal, 11pt, justify, indent 1 TAB)
    6. English title (normal, 12pt, UPPERCASE, center) + 2 blank lines
    7. "Abstract" label (bold, 11pt, indent 1 TAB)
    8. English abstract (italic, 11pt, justify, indent 1 TAB) + 4 blank lines
    9. "Bibliografie" (bold, 12pt, center) with 3 blank lines before
   10. Bibliography entries (justify, [nr.] format)
- Images: centered, caption normal Size 9, center
- Tables: centered, caption bold Size 10, center
"""

import sys
import re
import copy
import argparse
from pathlib import Path
from docx import Document
from docx.shared import Pt, Mm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree


# Constants
FONT_NAME = "Times New Roman"
FONT_SIZE_BODY = Pt(11)
FONT_SIZE_TITLE = Pt(12)
FONT_SIZE_IMG_CAPTION = Pt(9)
FONT_SIZE_TBL_CAPTION = Pt(10)
TAB_INDENT = Mm(12.7)
LINE_SPACING = 1.0

# Page setup: ISO B5
PAGE_WIDTH = Mm(176)
PAGE_HEIGHT = Mm(250.1)
MARGIN_TOP = Mm(20)
MARGIN_BOTTOM = Mm(20)
MARGIN_LEFT = Mm(20)
MARGIN_RIGHT = Mm(20)
HEADER_DISTANCE = Mm(12.7)
FOOTER_DISTANCE = Mm(12.7)


class DocumentSection:
    """Represents a semantic section of the document."""
    PREAMBLE = "preamble"       # stuff before title (headers, logos, etc.)
    TITLE_RO = "title_ro"       # Romanian title
    AUTHORS = "authors"         # Author names
    REZUMAT_LABEL = "rezumat_label"
    REZUMAT_TEXT = "rezumat_text"
    BODY = "body"               # Main body text
    TITLE_EN = "title_en"       # English title
    ABSTRACT_LABEL = "abstract_label"
    ABSTRACT_TEXT = "abstract_text"
    BIBLIOGRAPHY_HEADER = "bibliography_header"
    BIBLIOGRAPHY = "bibliography"
    POSTAMBLE = "postamble"     # stuff after bibliography


def detect_sections(doc):
    """
    Analyze the document and assign a semantic section to each paragraph.
    Returns a list of (paragraph_index, section_type) tuples.
    """
    paragraphs = doc.paragraphs
    n = len(paragraphs)
    sections = [None] * n

    # Find key markers
    rezumat_idx = None
    abstract_idx = None
    biblio_idx = None
    title_start = None
    title_end = None

    for i, p in enumerate(paragraphs):
        text = p.text.strip()
        text_upper = text.upper()
        style = p.style.name if p.style else ""

        # Find "REZUMAT" or "REZUMAT / ABSTRACT" heading
        if rezumat_idx is None and ("REZUMAT" in text_upper and len(text) < 80):
            rezumat_idx = i

        # Find bibliography section
        if biblio_idx is None and ("BIBLIOGRAFI" in text_upper or "REFERINȚE" in text_upper or "REFERINTE" in text_upper):
            if "Heading" in style or len(text) < 80:
                biblio_idx = i

    # Find the main title - look for the first substantial bold centered text
    # Skip empty lines and header/branding at the top
    for i, p in enumerate(paragraphs):
        text = p.text.strip()
        if not text:
            continue
        # Check if it looks like a title (centered, bold, substantial)
        is_centered = p.alignment == WD_ALIGN_PARAGRAPH.CENTER
        is_bold = any(r.font.bold for r in p.runs if r.font.bold is not None)
        has_heading_style = "Heading" in (p.style.name if p.style else "")

        # Skip small branding/header lines
        if is_centered and is_bold and len(text) > 15:
            # Check if this could be the title
            if title_start is None:
                # Skip lines that look like organization names
                if any(kw in text.lower() for kw in ["institute", "institut", "consortium", "association", "virtual"]):
                    continue
                if "LUCRARE" in text.upper() and "ȘTIINȚIFICĂ" in text.upper():
                    continue
                title_start = i
                title_end = i
            elif title_start is not None and title_end == i - 1:
                # Continuation of multi-line title
                if is_centered and is_bold:
                    title_end = i
            else:
                break
        elif title_start is not None:
            # Check if next non-empty is still title continuation
            if title_end == i - 1 and is_centered and is_bold:
                title_end = i
            elif text and title_end is not None:
                break

    # Now assign sections
    in_section = DocumentSection.PREAMBLE

    for i in range(n):
        text = paragraphs[i].text.strip()

        if title_start is not None and title_start <= i <= title_end:
            sections[i] = DocumentSection.TITLE_RO
            continue

        if i == rezumat_idx:
            sections[i] = DocumentSection.REZUMAT_LABEL
            in_section = "after_rezumat_label"
            continue

        if i == biblio_idx:
            sections[i] = DocumentSection.BIBLIOGRAPHY_HEADER
            in_section = DocumentSection.BIBLIOGRAPHY
            continue

        if in_section == DocumentSection.PREAMBLE:
            if title_start is not None and i < title_start:
                sections[i] = DocumentSection.PREAMBLE
            elif title_end is not None and i > title_end:
                # After title, before rezumat - these are authors/metadata
                if rezumat_idx is not None and i < rezumat_idx:
                    if text:
                        sections[i] = DocumentSection.AUTHORS
                    else:
                        sections[i] = DocumentSection.PREAMBLE
                else:
                    sections[i] = DocumentSection.BODY
            else:
                sections[i] = DocumentSection.PREAMBLE
            continue

        if in_section == "after_rezumat_label":
            if text and i < (biblio_idx or n):
                # The abstract text is typically the next non-empty paragraphs
                # Check if we've moved to body (heading or section number)
                style = paragraphs[i].style.name if paragraphs[i].style else ""
                if "Heading" in style or re.match(r'^\d+\.', text):
                    in_section = DocumentSection.BODY
                    sections[i] = DocumentSection.BODY
                else:
                    sections[i] = DocumentSection.REZUMAT_TEXT
            elif not text:
                # Empty line - could be transition to body
                # Look ahead to see what comes next
                next_text_idx = None
                for j in range(i + 1, min(i + 5, n)):
                    if paragraphs[j].text.strip():
                        next_text_idx = j
                        break
                if next_text_idx:
                    next_style = paragraphs[next_text_idx].style.name if paragraphs[next_text_idx].style else ""
                    next_text = paragraphs[next_text_idx].text.strip()
                    if "Heading" in next_style or re.match(r'^\d+\.', next_text):
                        in_section = DocumentSection.BODY
                        sections[i] = DocumentSection.BODY
                    else:
                        sections[i] = DocumentSection.REZUMAT_TEXT
                else:
                    sections[i] = DocumentSection.BODY
                    in_section = DocumentSection.BODY
            continue

        if in_section == DocumentSection.BODY:
            if biblio_idx is not None and i >= biblio_idx:
                sections[i] = DocumentSection.BIBLIOGRAPHY
            else:
                sections[i] = DocumentSection.BODY
            continue

        if in_section == DocumentSection.BIBLIOGRAPHY:
            if not text:
                # Check if we've reached the postamble (empty lines followed by centered text at end)
                remaining_non_empty = [j for j in range(i+1, n) if paragraphs[j].text.strip()]
                if remaining_non_empty:
                    all_centered = all(
                        paragraphs[j].alignment == WD_ALIGN_PARAGRAPH.CENTER
                        for j in remaining_non_empty
                    )
                    if all_centered:
                        in_section = DocumentSection.POSTAMBLE
                        sections[i] = DocumentSection.POSTAMBLE
                        continue
                sections[i] = DocumentSection.BIBLIOGRAPHY
            else:
                sections[i] = DocumentSection.BIBLIOGRAPHY
            continue

        if in_section == DocumentSection.POSTAMBLE:
            sections[i] = DocumentSection.POSTAMBLE
            continue

        sections[i] = in_section if in_section else DocumentSection.BODY

    # Fill any remaining None
    for i in range(n):
        if sections[i] is None:
            sections[i] = DocumentSection.BODY

    return sections, {
        "title_start": title_start,
        "title_end": title_end,
        "rezumat_idx": rezumat_idx,
        "biblio_idx": biblio_idx,
    }


def set_paragraph_format(paragraph, font_size=FONT_SIZE_BODY, bold=False, italic=False,
                         alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=None,
                         space_before=Pt(0), space_after=Pt(0), line_spacing=LINE_SPACING,
                         keep_content=True):
    """Apply formatting to a paragraph and all its runs."""
    pf = paragraph.paragraph_format
    pf.alignment = alignment
    pf.space_before = space_before
    pf.space_after = space_after
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.line_spacing = line_spacing

    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    else:
        pf.first_line_indent = None

    # Reset left/right indent
    pf.left_indent = None
    pf.right_indent = None

    for run in paragraph.runs:
        run.font.name = FONT_NAME
        run.font.size = font_size
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = RGBColor(0, 0, 0)

        # Set Times New Roman for East Asian text too
        rpr = run._element.get_or_add_rPr()
        rFonts = rpr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}" w:cs="{FONT_NAME}"/>')
            rpr.insert(0, rFonts)
        else:
            rFonts.set(qn('w:ascii'), FONT_NAME)
            rFonts.set(qn('w:hAnsi'), FONT_NAME)
            rFonts.set(qn('w:cs'), FONT_NAME)


def create_empty_paragraph(doc, before_element=None):
    """Create a properly formatted empty paragraph."""
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    return p


def setup_page(doc):
    """Configure page setup for all sections."""
    for section in doc.sections:
        section.page_width = PAGE_WIDTH
        section.page_height = PAGE_HEIGHT
        section.top_margin = MARGIN_TOP
        section.bottom_margin = MARGIN_BOTTOM
        section.left_margin = MARGIN_LEFT
        section.right_margin = MARGIN_RIGHT
        section.header_distance = HEADER_DISTANCE
        section.footer_distance = FOOTER_DISTANCE
        section.orientation = WD_ORIENT.PORTRAIT


def set_default_style(doc):
    """Set document default style."""
    try:
        style = doc.styles['Normal']
    except KeyError:
        # Some documents don't have 'Normal' - try default paragraph style
        from docx.enum.style import WD_STYLE_TYPE
        style = None
        for s in doc.styles:
            if s.type == WD_STYLE_TYPE.PARAGRAPH and s.name in ('Normal', 'Default', 'Standard'):
                style = s
                break
        if style is None:
            return
    font = style.font
    font.name = FONT_NAME
    font.size = FONT_SIZE_BODY
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.line_spacing = LINE_SPACING
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def format_document(input_path, output_path, author_names=None, title_en=None, abstract_en=None):
    """
    Main formatting function.

    Args:
        input_path: Path to input .docx
        output_path: Path for formatted output .docx
        author_names: Optional list of author names (if not auto-detected)
        title_en: Optional English title (if the document needs one added)
        abstract_en: Optional English abstract text
    """
    doc = Document(input_path)

    print(f"Loaded document with {len(doc.paragraphs)} paragraphs")

    # Detect sections
    sections, markers = detect_sections(doc)

    # Print detection results for debugging
    section_counts = {}
    for s in sections:
        section_counts[s] = section_counts.get(s, 0) + 1
    print(f"Detected sections: {section_counts}")
    print(f"Markers: {markers}")

    # Setup page
    setup_page(doc)
    set_default_style(doc)

    # Now we build a new document with proper structure
    new_doc = Document()
    setup_page(new_doc)
    set_default_style(new_doc)

    # Remove default empty paragraph
    if new_doc.paragraphs:
        p = new_doc.paragraphs[0]._element
        p.getparent().remove(p)

    # Collect content by section
    title_paragraphs = []
    author_paragraphs = []
    rezumat_text_paragraphs = []
    body_paragraphs = []
    biblio_header = None
    biblio_paragraphs = []

    for i, (para, sec) in enumerate(zip(doc.paragraphs, sections)):
        if sec == DocumentSection.TITLE_RO:
            title_paragraphs.append(para)
        elif sec == DocumentSection.AUTHORS:
            author_paragraphs.append(para)
        elif sec == DocumentSection.REZUMAT_LABEL:
            pass  # We'll create this ourselves
        elif sec == DocumentSection.REZUMAT_TEXT:
            if para.text.strip():
                rezumat_text_paragraphs.append(para)
        elif sec == DocumentSection.BODY:
            body_paragraphs.append(para)
        elif sec == DocumentSection.BIBLIOGRAPHY_HEADER:
            biblio_header = para
        elif sec == DocumentSection.BIBLIOGRAPHY:
            biblio_paragraphs.append(para)

    # === BUILD NEW DOCUMENT ===

    # 1. TITLE (bold, 12pt, UPPERCASE, center)
    title_text = " ".join(p.text.strip() for p in title_paragraphs if p.text.strip())
    if title_text:
        p = new_doc.add_paragraph()
        run = p.add_run(title_text.upper())
        set_paragraph_format(p, font_size=FONT_SIZE_TITLE, bold=True,
                           alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # 6 blank lines after title
    for _ in range(6):
        p = new_doc.add_paragraph()
        run = p.add_run("")
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT)

    # 2. AUTHORS (normal, 11pt, center)
    if author_names:
        for name in author_names:
            p = new_doc.add_paragraph()
            run = p.add_run(name)
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    elif author_paragraphs:
        for ap in author_paragraphs:
            if ap.text.strip():
                p = new_doc.add_paragraph()
                run = p.add_run(ap.text.strip())
                set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # 1 blank line after authors
    p = new_doc.add_paragraph()
    run = p.add_run("")
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT)

    # 3. "Rezumat" label (bold, 11pt, justify, indent 1 TAB)
    p = new_doc.add_paragraph()
    run = p.add_run("Rezumat")
    set_paragraph_format(p, bold=True, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                        first_line_indent=TAB_INDENT)

    # 4. Romanian abstract text (italic, 11pt, justify, indent 1 TAB)
    for rp in rezumat_text_paragraphs:
        if rp.text.strip():
            p = new_doc.add_paragraph()
            run = p.add_run(rp.text.strip())
            set_paragraph_format(p, italic=True, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                               first_line_indent=TAB_INDENT)

    # 2 blank lines after abstract
    for _ in range(2):
        p = new_doc.add_paragraph()
        run = p.add_run("")
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT)

    # 5. BODY TEXT
    for bp in body_paragraphs:
        text = bp.text.strip()
        style_name = bp.style.name if bp.style else ""
        is_heading1 = "Heading 1" in style_name
        is_heading2 = "Heading 2" in style_name
        is_heading = is_heading1 or is_heading2
        is_list = "List" in style_name

        if not text and not is_heading:
            # Preserve blank lines in body but don't add excessive ones
            p = new_doc.add_paragraph()
            run = p.add_run("")
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT)
            continue

        if not text:
            continue

        p = new_doc.add_paragraph()

        if is_heading:
            # Headings: bold, appropriate size, left-aligned or as body
            run = p.add_run(text)
            if is_heading1:
                set_paragraph_format(p, font_size=FONT_SIZE_BODY, bold=True,
                                   alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
            else:
                set_paragraph_format(p, font_size=FONT_SIZE_BODY, bold=True,
                                   alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                                   first_line_indent=TAB_INDENT)
        elif is_list:
            # List items - preserve with indent
            run = p.add_run(text)
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                               first_line_indent=TAB_INDENT)
        else:
            # Regular body text
            run = p.add_run(text)
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                               first_line_indent=TAB_INDENT)

    # 6-8. English title and abstract (if provided)
    if title_en:
        # 2 blank lines before English title
        for _ in range(2):
            p = new_doc.add_paragraph()
            run = p.add_run("")
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT)

        p = new_doc.add_paragraph()
        run = p.add_run(title_en.upper())
        set_paragraph_format(p, font_size=FONT_SIZE_TITLE,
                           alignment=WD_ALIGN_PARAGRAPH.CENTER)

        # 2 blank lines after English title
        for _ in range(2):
            p = new_doc.add_paragraph()
            run = p.add_run("")
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT)

        # "Abstract" label
        p = new_doc.add_paragraph()
        run = p.add_run("Abstract")
        set_paragraph_format(p, bold=True, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                           first_line_indent=TAB_INDENT)

        if abstract_en:
            p = new_doc.add_paragraph()
            run = p.add_run(abstract_en)
            set_paragraph_format(p, italic=True, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                               first_line_indent=TAB_INDENT)

        # 4 blank lines after English abstract
        for _ in range(4):
            p = new_doc.add_paragraph()
            run = p.add_run("")
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT)

    # 9. BIBLIOGRAPHY
    if biblio_paragraphs or biblio_header:
        # 3 blank lines before bibliography
        for _ in range(3):
            p = new_doc.add_paragraph()
            run = p.add_run("")
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT)

        # "Bibliografie" header (bold, 12pt, center)
        p = new_doc.add_paragraph()
        run = p.add_run("Bibliografie")
        set_paragraph_format(p, font_size=FONT_SIZE_TITLE, bold=True,
                           alignment=WD_ALIGN_PARAGRAPH.CENTER)

        # Bibliography entries
        ref_num = 1
        for bp in biblio_paragraphs:
            text = bp.text.strip()
            if not text:
                continue
            style_name = bp.style.name if bp.style else ""
            # Skip sub-headings within bibliography
            if "Heading" in style_name:
                continue

            # Format as [nr.] entry
            # Check if already has a number prefix
            if re.match(r'^\[?\d+[\].]', text):
                formatted = text
            else:
                formatted = f"[{ref_num}] {text}"
                ref_num += 1

            p = new_doc.add_paragraph()
            run = p.add_run(formatted)
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # Handle tables in the original document
    for table in doc.tables:
        # Copy tables (basic approach - preserves structure)
        new_table = new_doc.add_table(rows=len(table.rows), cols=len(table.columns))
        new_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                new_table.rows[i].cells[j].text = cell.text

    new_doc.save(output_path)
    print(f"\nFormatted document saved to: {output_path}")
    print(f"Total paragraphs in output: {len(new_doc.paragraphs)}")
    return new_doc


def main():
    parser = argparse.ArgumentParser(
        description="Format a .docx document according to conference rules"
    )
    parser.add_argument("input", help="Input .docx file path")
    parser.add_argument("-o", "--output", help="Output .docx file path (default: input_formatted.docx)")
    parser.add_argument("--authors", nargs="+", help="Author names (overrides auto-detection)")
    parser.add_argument("--title-en", help="English title")
    parser.add_argument("--abstract-en", help="English abstract text")

    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: File not found: {input_path}")
        sys.exit(1)

    if args.output:
        output_path = Path(args.output)
    else:
        output_path = input_path.with_stem(input_path.stem + "_formatted")

    format_document(
        str(input_path),
        str(output_path),
        author_names=args.authors,
        title_en=args.title_en,
        abstract_en=args.abstract_en,
    )


if __name__ == "__main__":
    main()
